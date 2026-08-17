#!/usr/bin/env python3
"""Generate the TGS-2024049781 Lesson Plan (LP) DOCX in the Tertiary house format.

Fast-Track Innovations with Agile Design Thinking and Generative AI (GenAI).

Cover page + Document Version Control Record + auto TOC + Arial 11pt body +
colour-coded 2-day schedule tables (9:30am-6:30pm, 8 training hours/day, 1h
lunch, tea breaks counted within training time, final assessment on Day 2).
Topics and activities come from course_data + the domain data files so the LP
stays aligned with the deck, the Learner Guide and the activities/ folder.
"""
import os, sys
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

HERE = os.path.dirname(os.path.abspath(__file__)); sys.path.insert(0, HERE)
import course_data as C
from data_domain1 import DOMAIN1
from data_domain2 import DOMAIN2
from data_domain3 import DOMAIN3
from data_domain4 import DOMAIN4
ACT = DOMAIN1 + DOMAIN2 + DOMAIN3 + DOMAIN4
import prodoc


def _find_repo(start):
    env = os.environ.get("COURSE_REPO")
    if env and os.path.isdir(env):
        return env
    d = start
    for _ in range(8):
        d = os.path.dirname(d)
        if os.path.isdir(os.path.join(d, "courseware")) and os.path.isdir(os.path.join(d, "labs")):
            return d
    return os.path.dirname(os.path.dirname(HERE))


REPO = _find_repo(HERE)
ASSETS = os.path.join(os.path.dirname(HERE), "assets")

BRAND = RGBColor(0x1F, 0x6F, 0xEB); DARK = RGBColor(0x11, 0x18, 0x27); GREY = RGBColor(0x55, 0x5B, 0x66)
HEADER_FILL = "1F6FEB"; TOPIC_FILL = "E8F0FE"; BREAK_FILL = "FFF4E5"
LUNCH_FILL = "FDE9D9"; ASSESS_FILL = "E8F7EE"


def act_titles(nums):
    return "; ".join(f"Activity {a['num']}: {a['title']}" for a in ACT if a['num'] in nums)


# ------------------------------------------------ schedule (single source of truth for timing)
# (start, end, minutes, kind, activity_text)  kind: admin/topic/activity/break/lunch/assess/recap
# Day totals must be 480 training minutes (the 1-hour lunch is excluded).
SCHEDULE = {
 1: (C.DAY_THEMES[1], [
    ("9:30", "10:00", 30, "admin",
     "Welcome, trainer and learner introductions, ground rules, learning outcomes, course outline and mandatory SSG digital attendance (AM). Introduction to the two course ed-tools: Design Thinking Studio and the Padlet Classroom Board."),
    ("10:00", "11:00", 60, "topic",
     "Topic 1 — Foundations of Design Thinking, Agile and Generative AI: what design thinking is; the three mindsets (Design Thinking, Lean, Agile) and the question each answers; the Double Diamond; the five stages; wicked problems."),
    ("11:00", "11:15", 15, "break", "Tea break"),
    ("11:15", "12:00", 45, "topic",
     "Topic 1 continued — Product development through design thinking and agile; how generative AI enhances the loop; where AI must not replace human judgement; latest trends in innovation."),
    ("12:00", "13:00", 60, "activity",
     "Hands-on case study: " + act_titles([1])),
    ("13:00", "14:00", 60, "lunch", "Lunch break"),
    ("14:00", "14:15", 15, "admin", "Mandatory SSG digital attendance (PM)"),
    ("14:15", "15:30", 75, "activity",
     "Hands-on case studies: " + act_titles([2, 3])),
    ("15:30", "15:45", 15, "break", "Tea break"),
    ("15:45", "16:45", 60, "topic",
     "Topic 2 — Problem Framing and Ideation: framing the design challenge; the Problem-Assumption Model; problem statement vs Point of View vs How Might We; the altitude test; empathy maps and persona mapping with GenAI; ideation techniques."),
    ("16:45", "18:15", 90, "activity",
     "Hands-on case studies: " + act_titles([4, 5])),
    ("18:15", "18:30", 15, "recap", "Day 1 recap, Q&A and close"),
 ]),
 2: (C.DAY_THEMES[2], [
    ("9:30", "9:45", 15, "recap", "Day 1 recap and mandatory SSG digital attendance (AM)"),
    ("9:45", "10:45", 60, "activity",
     "Topic 2 concluded — prototyping with AI-driven tools and defining a falsifiable test plan. Hands-on case study: " + act_titles([6])),
    ("10:45", "11:00", 15, "break", "Tea break"),
    ("11:00", "12:00", 60, "topic",
     "Topic 3 — Agile Development and AI for Rapid Solution Delivery: being agile vs doing agile; the Scrum loop and the three accountabilities; dual-track agile; the mini-waterfall anti-pattern; Epic → Feature → Hill → Story → Task."),
    ("12:00", "13:00", 60, "activity",
     "Hands-on case study: " + act_titles([7])),
    ("13:00", "14:00", 60, "lunch", "Lunch break"),
    ("14:00", "14:15", 15, "admin", "Mandatory SSG digital attendance (PM)"),
    ("14:15", "15:30", 75, "activity",
     "Topic 3 continued — user stories and the 3 C's, the three-cake release model, story points and Planning Poker, GenAI across the sprint. Hands-on case studies: " + act_titles([8, 9])),
    ("15:30", "15:45", 15, "break", "Tea break"),
    ("15:45", "16:45", 60, "topic",
     "Topic 4 — Scaling and Sustaining Innovations: the three levers that scale innovation; culture and the four principles; stakeholder buy-in; resource management with AI; sensemaking; systems thinking and feedback loops; metrics and KPIs."),
    ("16:45", "17:45", 60, "activity",
     "Hands-on case studies: " + act_titles([10, 11, 12])),
    ("17:45", "18:00", 15, "recap", "Course revision and Q&A"),
    ("18:00", "18:30", 30, "assess",
     "Briefing for Assessment, mandatory TRAQOM survey and Assessment digital attendance. End of the 8 instructional hours."),
 ]),
}

# ------------------------------------------------ build document
doc = Document()
normal = doc.styles["Normal"]; normal.font.name = "Arial"; normal.font.size = Pt(11)
prodoc.style_headings(doc)

prodoc.add_cover_page(doc, "LESSON PLAN", C.TITLE, C.VERSION.lstrip("v"),
                      org_logo=os.path.join(ASSETS, "tertiary-infotech-logo.png"),
                      course_logo=None, course_code=C.COURSE_CODE)
prodoc.add_version_control(doc, [
 ("3.0", "1 January 2026",
  "Legacy master trainer deck and lesson plan (v3) — 4 topics, exercise-based delivery.",
  C.TRAINER),
 ("4.0", C.VERSION_DATE,
  "Major revision. Content beefed up from Thoughtworks, Adaptovate, BMC, PMI, UX Magazine, "
  "The Design Gym, PremierAgile and Startup Frontier. Twelve real-world case-study activities "
  "(Airbnb, Netflix/Blockbuster, GE Adventure Series, IDEO, Spotify, IBM, DBS, Boeing 737 MAX) "
  "replace the generic exercises, each with scenario, discussion questions and debrief. "
  "Design Thinking Studio and Padlet Classroom Board adopted as the class ed-tools.",
  C.TRAINER),
])
prodoc.add_toc(doc)


def H(text, level=1):
    return doc.add_heading(text, level=level)


H("Course Information", 1)
info = [("Course Title", C.TITLE),
        ("WSQ Course Reference", C.COURSE_CODE),
        ("Training Provider", C.ORG + "  (" + C.UEN.replace('UEN: ', 'UEN ') + ")"),
        ("Duration", "2 days · 8 training hours per day (16 instructional hours), plus a separate 2-hour assessment session"),
        ("Daily Timing", "9:30 am – 6:30 pm (1-hour lunch; tea breaks counted within training time)"),
        ("Mode", "Instructor-led, case-study based, with collaborative online ed-tools"),
        ("Skills Framework", f"{C.TSC_TITLE} ({C.TSC_CODE})"),
        ("Trainer", C.TRAINER)]
t = doc.add_table(rows=0, cols=2); t.style = "Table Grid"
for k, v in info:
    c = t.add_row().cells
    c[0].text = ""; r = c[0].paragraphs[0].add_run(k); r.bold = True; r.font.size = Pt(10)
    prodoc._shade_cell(c[0], TOPIC_FILL)
    c[1].text = ""; c[1].paragraphs[0].add_run(v).font.size = Pt(10)

H("Learning Outcomes", 1)
doc.add_paragraph("On completion of this course, learners will be able to:")
for lo in C.LEARNING_OUTCOMES:
    p = doc.add_paragraph(style="List Bullet"); p.add_run(lo).font.size = Pt(10.5)

H("Skills Framework Alignment", 1)
doc.add_paragraph(f"TSC Title: {C.TSC_TITLE}   ·   TSC Code: {C.TSC_CODE}")
doc.add_paragraph("Abilities:")
for a in C.TSC_ABILITIES:
    p = doc.add_paragraph(style="List Bullet"); p.add_run(a).font.size = Pt(10)
doc.add_paragraph("Knowledge:")
for k in C.TSC_KNOWLEDGE:
    p = doc.add_paragraph(style="List Bullet"); p.add_run(k).font.size = Pt(10)

H("Teaching and Learning Methods", 1)
for m in [
    "Trainer-led concept delivery using a highly visual slide deck (diagrams, comparison tables and frameworks rather than bullet walls).",
    "Twelve real-world case-study activities, each carrying a documented case, a workplace scenario, five discussion questions and a trainer debrief.",
    "Collaborative group work in groups of 3–5 using two online ed-tools: the Design Thinking Studio (5-stage workspace) and the Padlet Classroom Board.",
    "Guided use of generative AI tools (ChatGPT / Microsoft Copilot) with explicit critique of AI output against evidence.",
    "Cross-group peer review — groups critique each other's output and defend their reasoning.",
    "Continuous formative feedback from the trainer during each activity, and a structured debrief after each one.",
]:
    p = doc.add_paragraph(style="List Bullet"); p.add_run(m).font.size = Pt(10.5)

H("Learning Resources", 1)
for m in [
    "Trainer Slides (PPT/PDF), Learner Guide (PDF) and this Lesson Plan (PDF) — downloadable from the LMS at https://lms-tms.tertiaryinfotech.com.",
    "Twelve individual activity briefs (PDF), one per activity, each containing the case, the scenario, the discussion questions, the step-by-step instructions and the debrief.",
    "Design Thinking Studio — https://alfredang.github.io/designthinking/",
    "Padlet Classroom Board — https://alfredang.github.io/padlet/",
    "A generative AI assistant (ChatGPT or Microsoft Copilot) and a laptop with a modern web browser.",
]:
    p = doc.add_paragraph(style="List Bullet"); p.add_run(m).font.size = Pt(10.5)

H("Assessment", 1)
for a in [C.ASSESSMENT["written"], C.ASSESSMENT["practical"],
          "Format: Open Book — course slides, Learner Guide, activity briefs and approved materials only.",
          "The final assessment is conducted as a separate session at the end of Day 2, outside the 16 instructional hours: WA (0.5 h) then CS (1 h).",
          "Learners are assessed as Competent (C) or Not Yet Competent (NYC) on each instrument.",
          C.ASSESSMENT["note"]]:
    p = doc.add_paragraph(style="List Bullet"); p.add_run(a).font.size = Pt(10.5)


def set_cell(cell, text, bold=False, size=9.5, color=None, fill=None, align=None):
    cell.text = ""; p = cell.paragraphs[0]
    if align:
        p.alignment = align
    r = p.add_run(text); r.bold = bold; r.font.size = Pt(size); r.font.name = "Arial"
    if color:
        r.font.color.rgb = color
    if fill:
        prodoc._shade_cell(cell, fill)


KIND_FILL = {"topic": TOPIC_FILL, "break": BREAK_FILL, "lunch": LUNCH_FILL,
             "assess": ASSESS_FILL, "admin": "F3F5F8", "recap": "F3F5F8", "activity": None}

H("Course Schedule", 1)
for day, (theme, rows) in SCHEDULE.items():
    H(f"Day {day} — {theme}", 2)
    tbl = doc.add_table(rows=0, cols=3); tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = tbl.add_row().cells
    for i, htext in enumerate(["Time", "Duration", "Topic / Activity"]):
        set_cell(hdr[i], htext, bold=True, size=10, color=RGBColor(0xFF, 0xFF, 0xFF), fill=HEADER_FILL)
    training = 0
    for start, end, mins, kind, text in rows:
        cells = tbl.add_row().cells; fill = KIND_FILL.get(kind)
        set_cell(cells[0], f"{start}–{end}", bold=(kind in ("topic", "assess")), size=9.5, fill=fill)
        set_cell(cells[1], f"{mins} min", size=9.5, fill=fill)
        set_cell(cells[2], text, bold=(kind in ("topic", "assess")), size=9.5, fill=fill)
        if kind != "lunch":
            training += mins
    for row in tbl.rows:
        row.cells[0].width = Inches(1.15); row.cells[1].width = Inches(0.9); row.cells[2].width = Inches(4.75)
    p = doc.add_paragraph()
    r = p.add_run(f"Total training time: {training} minutes ({training // 60} hours).")
    r.italic = True; r.font.size = Pt(9.5); r.font.color.rgb = GREY
    assert training == 480, f"Day {day} training minutes = {training}, expected 480"

H("Assessment Session (conducted separately, after the 16 instructional hours)", 1)
doc.add_paragraph(
    "The 2-hour assessment is scheduled as a separate session at the end of Day 2 and is NOT counted "
    "within the 16 instructional hours. Learners who have met the 75% attendance requirement sit both "
    "instruments in the order below.")
ast_ = doc.add_table(rows=0, cols=3); ast_.style = "Table Grid"
hdr = ast_.add_row().cells
for i, htext in enumerate(["Duration", "Instrument", "Details"]):
    set_cell(hdr[i], htext, bold=True, size=10, color=RGBColor(0xFF, 0xFF, 0xFF), fill=HEADER_FILL)
for dur, inst, det in [
    ("15 min", "Briefing for Assessment",
     "Assessment rules, materials permitted, and the appeal process."),
    ("0.5 hour", "Written Assessment (WA) — Short-Answer Questions (SAQ)",
     "6 open-ended knowledge questions covering K1–K5. Open book."),
    ("1 hour", "Case Study (CS)",
     "1 scenario with 4 tasks covering A1–A7. Open book."),
    ("15 min", "Submission and records",
     "Learners upload answers to the LMS and sign the Assessment Summary Record."),
]:
    cells = ast_.add_row().cells
    set_cell(cells[0], dur, bold=True, size=9.5, fill=ASSESS_FILL)
    set_cell(cells[1], inst, bold=True, size=9.5, fill=ASSESS_FILL)
    set_cell(cells[2], det, size=9.5)
for row in ast_.rows:
    row.cells[0].width = Inches(0.9); row.cells[1].width = Inches(2.6); row.cells[2].width = Inches(3.3)
doc.add_paragraph("")

H("Activity Reference (aligned to topics and learning outcomes)", 1)
tt = doc.add_table(rows=0, cols=4); tt.style = "Table Grid"
hdr = tt.add_row().cells
for i, htext in enumerate(["Topic", "Weighting", "Activities", "Ed-tool"]):
    set_cell(hdr[i], htext, bold=True, size=10, color=RGBColor(0xFF, 0xFF, 0xFF), fill=HEADER_FILL)
for tp in C.TOPICS:
    acts = [a for a in ACT if a["topic"] == tp["num"]]
    cells = tt.add_row().cells
    set_cell(cells[0], f"Topic {tp['code']}: {tp['title']}", bold=True, size=9.5, fill=TOPIC_FILL)
    set_cell(cells[1], tp["weighting"], size=9.5, fill=TOPIC_FILL)
    set_cell(cells[2], ", ".join(f"Activity {a['num']}" for a in acts), size=9.5)
    tools = set()
    for a in acts:
        if "Design Thinking Studio" in a["services"]:
            tools.add("Design Thinking Studio")
        if "Padlet" in a["services"]:
            tools.add("Padlet")
    set_cell(cells[3], " / ".join(sorted(tools)) or "—", size=9.5)
for row in tt.rows:
    row.cells[0].width = Inches(3.0); row.cells[1].width = Inches(0.8)
    row.cells[2].width = Inches(1.9); row.cells[3].width = Inches(1.4)

H("Detailed Activity Schedule", 1)
at = doc.add_table(rows=0, cols=4); at.style = "Table Grid"
hdr = at.add_row().cells
for i, htext in enumerate(["#", "Activity", "Duration", "Real-world case"]):
    set_cell(hdr[i], htext, bold=True, size=10, color=RGBColor(0xFF, 0xFF, 0xFF), fill=HEADER_FILL)
for a in ACT:
    cells = at.add_row().cells
    set_cell(cells[0], str(a["num"]), bold=True, size=9.5, fill=TOPIC_FILL)
    set_cell(cells[1], a["title"], size=9.5)
    set_cell(cells[2], f"{a.get('duration', 45)} min", size=9.5)
    set_cell(cells[3], a["case"].split(".")[0][:90] + "…", size=9)
for row in at.rows:
    row.cells[0].width = Inches(0.4); row.cells[1].width = Inches(2.9)
    row.cells[2].width = Inches(0.8); row.cells[3].width = Inches(3.0)

prodoc.add_page_numbers(doc)
prodoc.enable_update_fields(doc)
OUT = os.path.join(REPO, "courseware", f"LP-{C.SHORT_TITLE}.docx")
doc.save(OUT)
print("Saved", OUT)
