#!/usr/bin/env python3
"""Generate the TGS-2024049781 Learner Guide as BOTH a Markdown mirror (LG-*.md at repo root)
and a DOCX (courseware/LG-*.docx) from one source, so they never diverge.

Fast-Track Innovations with Agile Design Thinking and Generative AI (GenAI).

House format: cover page, Document Version Control Record, auto TOC, Arial 11pt body,
full teaching content per topic, then one section per activity carrying the real-world
case, the scenario, the DETAILED STEP-BY-STEP instructions (which appear ONLY here, never
on the slides), the discussion questions and the trainer debrief. All content is driven by
course_data + the domain data files, keeping the LG 100% aligned with the deck, the Lesson
Plan and the activities/ folder.
"""
import os, sys
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

HERE = os.path.dirname(os.path.abspath(__file__)); sys.path.insert(0, HERE)
import course_data as C
from data_domain1 import DOMAIN1
from data_domain2 import DOMAIN2
from data_domain3 import DOMAIN3
from data_domain4 import DOMAIN4
ACT = DOMAIN1 + DOMAIN2 + DOMAIN3 + DOMAIN4
import prodoc


def _find_repo(start):
    env = os.environ.get("COURSE_REPO")
    if env and os.path.isdir(env):
        return env
    d = start
    for _ in range(8):
        d = os.path.dirname(d)
        if os.path.isdir(os.path.join(d, "courseware")) and os.path.isdir(os.path.join(d, "labs")):
            return d
    return os.path.dirname(os.path.dirname(HERE))


REPO = _find_repo(HERE)
ASSETS = os.path.join(os.path.dirname(HERE), "assets")
IMGDIR = os.path.join(ASSETS, "img")

# ---------------- block DSL (single content stream → MD + DOCX) ----------------
B = []
def h1(t): B.append(("h1", t))
def h2(t): B.append(("h2", t))
def h3(t): B.append(("h3", t))
def p(t):  B.append(("p", t))
def bullets(xs): B.append(("bullets", xs))
def numbered(xs): B.append(("numbered", xs))
def steps(xs): B.append(("steps", xs))
def note(t): B.append(("note", t))
def quote(t): B.append(("quote", t))
def table(headers, rows): B.append(("table", headers, rows))
def img(name, caption=""): B.append(("img", name, caption))
def rule(): B.append(("rule",))
def dl(pairs): B.append(("dl", pairs))


# ================================================================= INTRODUCTION
h1("Introduction")
p(f"This Learner Guide accompanies the WSQ course {C.TITLE} ({C.COURSE_CODE}), conducted by "
  f"{C.ORG} ({C.UEN.replace('UEN: ', 'UEN ')}). It is your reference during the two training days "
  "and your working manual afterwards.")
p("The guide is organised into four topics that match the slide deck and the Lesson Plan exactly. "
  "Each topic opens with the teaching content — the frameworks, comparisons and models you will be "
  "assessed on — and then works through its real-world case-study activities. For every activity "
  "you will find the documented case, the workplace scenario your group works on, the full "
  "step-by-step instructions, the discussion questions, and the debrief that states what the room "
  "should conclude.")
note("The detailed step-by-step instructions appear in this guide and in the individual activity "
     "briefs — deliberately NOT on the slides. The slides carry the case, the scenario and the "
     "questions so that the class discusses rather than reads. Follow the steps from here.")
p("This is an open-book course. You may use this guide, the slides, your activity briefs and any "
  "approved material during the final assessment.")

h1("Course Learning Outcomes")
p("On completion of this course you will be able to:")
bullets(C.LEARNING_OUTCOMES)

h1("Skills Framework Alignment")
p(f"This course is aligned to the Skills Framework Technical Skill and Competency "
  f"{C.TSC_TITLE} ({C.TSC_CODE}).")
h3("Abilities")
bullets(C.TSC_ABILITIES)
h3("Knowledge")
bullets(C.TSC_KNOWLEDGE)

# ================================================================= ED-TOOLS
h1("Before You Start — Your Course Tools")
p("Two collaborative web tools are used throughout this course. Both run in a browser, need no "
  "installation and no account, and are shared by the whole class. Bring a laptop; a phone works "
  "but is cramped for the workspace tool.")

h2("Design Thinking Studio")
p("Design Thinking Studio is a shared five-stage workspace at "
  "https://alfredang.github.io/designthinking/. The whole class works inside one live workspace and "
  "posts notes into structured sections under each of the five design thinking stages.")
h3("How to join a workspace")
steps([
    ("Open https://alfredang.github.io/designthinking/ in your browser.", ""),
    ("The trainer clicks Create New Workspace and shares the workspace code, the link, or a QR code (via the Share button).", ""),
    ("Enter the workspace code in the 'Enter workspace code' field and click Join Workspace.", ""),
    ("Enter your display name when prompted, then click Continue. Your name appears on every note you post.", ""),
    ("Use the stage stepper across the top to move between Empathize, Define, Ideate, Prototype and Test.", ""),
    ("Inside any section click Add note, then fill in the Title, the Description and an optional Category, and click Save.", ""),
])
h3("What is in each stage")
table(["Stage", "Sections available"],
      [["1. Empathize", "User Personas · User Pain Points · Interview Notes · Observation Notes · Empathy Map (Says / Thinks / Does / Feels)"],
       ["2. Define", "Problem Statement (voting) · How Might We Questions (voting) · Key Insights · User Needs · Success Criteria"],
       ["3. Ideate", "Brainstorming Board (voting) · Crazy 8s Ideas (voting) · Solution Sketches · Idea Categories · Voting & Prioritisation"],
       ["4. Prototype", "Prototype Description · Screens / Wireframes · Feature List · User Flow · Assumptions · Prototype Checklist"],
       ["5. Test", "Test Plan · User Feedback · Test Results · Issues Found · Improvements · Final Recommendations"]])
note("Sections marked with voting let the group mark notes as important — use this to converge. "
     "In the Test stage, Issues Found and Improvements can be marked resolved as you work through them.")

h2("Padlet Classroom Board")
p("The Padlet Classroom Board at https://alfredang.github.io/padlet/ is an online classroom board "
  "for posting group output, comparing it across groups and voting on the strongest thinking.")
h3("How to join a classroom")
steps([
    ("Open https://alfredang.github.io/padlet/ in your browser.", ""),
    ("The trainer creates the classroom and shares the classroom code.", ""),
    ("Click Join an existing classroom, enter the classroom code and click Join classroom.", ""),
    ("Enter your display name so your posts are attributed to you.", ""),
    ("Find the section for the current activity, then click + / Add post here.", ""),
    ("Give the post a Subject and body. You may attach an image, a link, a YouTube video or a PDF.", ""),
    ("Use Like on the strongest posts from other groups, and Comments to challenge their reasoning.", ""),
])
note("Sort the board by Newest first, Oldest first or Most liked. The trainer may pin an "
     "Announcement to the top of the board with the current task.")

h2("Your Generative AI Assistant")
p("You will use a generative AI assistant (ChatGPT, Microsoft Copilot or an equivalent) in most "
  "activities. Any current model is fine. Three working rules apply throughout this course:")
bullets([
    "Generate manually FIRST, then use AI to extend. Running AI first anchors the whole group to its framing and measurably narrows the range of ideas in the room.",
    "Treat every AI output as a hypothesis, never as evidence. Label AI-generated personas, insights and metrics as unvalidated until real-world evidence confirms them.",
    "Never paste confidential, personal or client-identifying data into a public AI tool. Use anonymised or synthetic details in class.",
])

# ================================================================= TOPIC 1
rule()
h1("Topic 1 — Foundations of Design Thinking, Agile, and Generative AI")
p("Topic weighting: 25%. This topic establishes the vocabulary for the whole course: what design "
  "thinking actually is, how it differs from lean and agile, and where generative AI genuinely "
  "helps rather than merely appearing to.")

h2("What Design Thinking Is")
p("Design Thinking is a solution-based approach to problem solving that is especially useful for "
  "problems that are ill-defined, ambiguous or simply unknown. It works by understanding the human "
  "needs involved, re-framing the problem in human-centred terms, generating many ideas in "
  "structured ideation sessions, and adopting a hands-on approach through prototyping and testing.")
p("The central discipline is resisting your first plausible answer. When was the last time your "
  "first idea was your best idea? Everyone designs, consciously or not — if you are solving a "
  "problem, you are designing a solution. Design thinking is a mindset that helps you do it better.")
h3("The abilities that make a designer")
p("Carissa Carter, head of teaching at the Stanford d.school, describes the abilities that "
  "distinguish strong designers:")
dl([("Dealing with ambiguity", "Staying productive when the problem itself is still unclear. This is the defining skill of the discipline and the hardest to fake."),
    ("Empathetic learning", "Understanding a user's world from inside their experience, rather than from a survey summary or a dashboard."),
    ("Synthesis", "Turning a large mass of qualitative evidence into a small number of insights you can actually design against."),
    ("Experimentation", "Treating every belief as testable, and designing the cheapest test that could disprove it.")])

h2("The Three Mindsets — Design Thinking, Lean and Agile")
img("three-mindsets.png", "The three mindsets of product development and where they overlap.")
p("The clearest formulation of how these three fit together comes from Thoughtworks:")
quote("Design Thinking is how we explore and solve problems. Lean is our framework for testing our "
      "beliefs and learning our way to the right outcomes. Agile is how we adapt to changing "
      "conditions with software.")
p("Compressed for memory: Design Thinking finds the RIGHT PROBLEM. Lean validates the RIGHT "
  "SOLUTION. Agile makes sure you BUILD IT RIGHT. They are completing, not competing — the "
  "question is never 'Lean or Agile?', it is 'and'.")
table(["", "Design Thinking", "Lean Startup", "Agile"],
      [["Primary focus", "Understanding users deeply", "Validating ideas efficiently", "Iterative delivery and adaptation"],
       ["The question it answers", "How should we think about this problem?", "How do we validate this solution efficiently?", "How do we build, scale and improve it?"],
       ["Origin", "Human-centred design practice", "Toyota lean manufacturing", "A counterpoint to Waterfall"],
       ["Core move", "Reframe the problem through empathy", "Build an MVP; let the customer determine value", "Ship increments in short sprints"],
       ["Timeframe", "Flexible — one day to one month", "Rapid, minimal validation cycles", "Sprint-based and continuous"],
       ["Fails when", "It never reaches delivery", "It validates a problem nobody has", "It builds the wrong thing efficiently"]])
note("A critical caution from the same source: all three mindsets are commonly corrupted by being "
     "codified into rituals and certifications and rolled out mindlessly. They are mindsets, not "
     "processes.")

h3("Problem finding versus problem solving")
p("The crispest two-word contrast: Design Thinking is problem FINDING; Agile is problem SOLVING. "
  "Agile focuses on solving a predefined issue efficiently, and can become dysfunctional when user "
  "engagement is minimal. Design Thinking focuses on selecting the right issue to address in the "
  "first place, and only begins once you understand user needs.")
p("PMI puts the same point in terms of destination: agile's iterative approach lets teams react to "
  "change and deliver finished products faster — but this is only valuable if it is the right "
  "destination to begin with.")

h2("The Double Diamond")
img("double-diamond.png", "The Double Diamond — the problem space and the solution space, each with a diverge and a converge phase.")
p("The Double Diamond describes innovation as two linked diamonds, each alternating divergent and "
  "convergent thinking.")
dl([("Diamond 1 — the problem space", "Diverge by zooming out to understand the customer, their context and their needs. Converge by formulating a clear problem statement. This diamond answers: are we solving the right problem?"),
    ("Diamond 2 — the solution space", "Diverge by ideating and testing multiple solutions. Converge by selecting the optimal one. This diamond answers: are we building the right solution?"),
    ("Then agile delivery", "Sprints answer the third question: are we building it right?")])
note("The first diamond is the one teams skip under deadline pressure. Skipping it does not save "
     "time — it relocates the cost to the point where the wrong thing has already been built.")

h2("The Five Stages of Design Thinking")
img("dt-five-stages.png", "The five stages — a loop you re-enter, not a waterfall you complete once.")
p("The five stages overlap and repeat continuously, with forward momentum on problem-solving always "
  "in play. The process is flexible and fluid by nature; a result at Test can send you straight back "
  "to Empathise.")
h3("1. Empathise")
p("Understand the experience, situation and emotions of the person you are designing for. Observe "
  "users and their behaviour in the context of their lives. Engage people in conversation and ask "
  "why. Watch and listen: ask someone to complete a task and narrate what they are doing. Where "
  "possible, immerse yourself in the physical environment to gain a deeper personal understanding.")
h3("2. Define")
p("Process and synthesise your findings to form a user Point of View that you will address. "
  "Develop an understanding of the type of person you are designing for, select a limited set of "
  "needs you think are important to fulfil, and express the insights you developed. Analyse your "
  "observations and synthesise them to define the core problem in a human-centred manner.")
h3("3. Ideate")
p("Focus on idea generation. Translate problems into solutions. Explore a wide variety and a large "
  "quantity of ideas so you can go beyond the obvious. Combine conscious and unconscious thought "
  "with rational analysis and imagination. Leverage the group to reach new ideas and build on other "
  "people's. Critically: separate the generation of ideas from their evaluation, to give "
  "imagination a voice.")
h3("4. Prototype")
p("Build to think. Create a simple, cheap and fast artefact that shapes the idea so you can "
  "experience and interact with it. Create something of deliberately low resolution — a physical "
  "object, a sketch, a paper screen. Quick and dirty is the point. Storyboard a scenario you can "
  "role-play so people can experience the solution.")
h3("5. Test")
p("Ask for feedback. Let people use the prototype hands-on and listen to what they say. Let users "
  "talk and describe how they feel. Learn about the user, reframe your view and refine the "
  "prototype. Solutions are accepted, improved and re-examined, or rejected, on the basis of the "
  "users' experiences. Fail sooner rather than later.")

h2("Wicked Problems")
p("The term wicked problem was coined in 1973 by Horst Rittel and Melvin Webber, professors of "
  "design and urban planning at UC Berkeley. Wicked problems are abstract and difficult to define, "
  "with multiple interconnected layers — poverty, climate change, public health. Their defining "
  "property is that they have no fixed endpoint and no single correct solution, only better and "
  "worse responses.")
p("A tame problem, by contrast, has a knowable answer. Reaching for a tame method on a wicked "
  "problem is one of the most common and expensive errors in organisational innovation, and it is "
  "precisely why framing must precede solving.")

h2("How Generative AI Enhances the Loop")
img("genai-overlay.png", "Where generative AI accelerates each stage of the design thinking loop.")
p("Generative AI compresses the expensive parts of the innovation loop. A seven-year study of design "
  "thinking by Jeanne Liedtka at the University of Virginia's Darden School found that immersion in "
  "user experience provides rich raw material for insight, but that finding patterns and making "
  "sense of the resulting mass of qualitative data is a significant challenge. That documented "
  "bottleneck — sense-making across large volumes of qualitative evidence — is exactly what "
  "generative AI relieves.")
table(["Stage", "GenAI genuinely accelerates", "What must stay human"],
      [["Empathise", "Synthesising interview transcripts and field notes at speed", "Being in the room; noticing what is not said"],
       ["Define", "Generating and stress-testing candidate reframes", "Choosing which problem the organisation will own"],
       ["Ideate", "Multiplying idea volume and range on demand", "Recognising the idea that fits this context"],
       ["Prototype", "Drafting copy, flows, screens and edge cases", "Deciding what is cheap enough to learn from"],
       ["Test", "Clustering feedback and drafting test scripts", "Watching a real user hesitate, and asking why"]])
quote("AI acts as an amplifier of good Agile and Design Thinking discipline rather than a substitute "
      "for it. The quality of outcomes still depends on the rigor of the underlying process and the "
      "judgment of the people running it.")
note("A language model generates the statistically likely, not the locally true. That is what you "
     "want for divergence, and what you must not trust for convergence.")

h2("Latest Trends in Innovation")
dl([("AI-assisted discovery", "Research synthesis that took a team weeks now takes hours; the bottleneck moves from analysis to judgement."),
    ("Continuous discovery", "Discovery is no longer a phase that precedes delivery — it runs permanently alongside it."),
    ("Cross-functional trios", "A product manager, a designer and an engineer own discovery together rather than passing artefacts down a chain."),
    ("Evidence over opinion", "Decisions increasingly require a named assumption and the test that would disprove it."),
    ("Responsible AI", "Provenance, bias and governance become part of the innovation process rather than a legal afterthought."),
    ("Outcome over output", "Organisations shift from counting features shipped to measuring the customer value created.")])

# ================================================================= TOPIC 2
rule()
h1("Topic 2 — Problem Framing and Ideation: Leveraging Design Thinking and Generative AI")
p("Topic weighting: 25%. This topic covers the highest-leverage move in the whole discipline — "
  "turning a technical brief into a human problem — and then the disciplined generation and "
  "convergence of ideas against it.")

h2("Framing the Design Challenge")
p("Human-centred design, as IDEO describes it, is an approach to problem-solving based on "
  "techniques that communicate with, interact with, empathise with and stimulate the people "
  "involved, in order to understand their needs, desires and experiences.")
p("Framing is about turning problems into opportunities and organising the possible solutions. At "
  "moments of ambiguity it clarifies where you should push your design. Follow six steps:")
numbered([
    "Write your design challenge — one short, memorable sentence conveying the problem you want to solve. For example: make families spend more time together.",
    "Frame it as a question, specifically a 'How Might We' question. For example: How might we help families spend more time together?",
    "Define the impact you would like to have in light of that challenge.",
    "Think about possible solutions to the problem — without committing to any of them yet.",
    "Write down the context and the constraints that surround the original question.",
    "Look back at your How Might We question and revise it. It is very common to adapt it as you dig deeper.",
])

h3("The altitude test")
p("The most common mistake is a How Might We question that is either too broad or too narrow.")
bullets([
    "TOO NARROW: you can immediately name only one obvious solution — the answer is hiding inside the question. 'How might we send an SMS reminder?' is a solution wearing a question mark.",
    "TOO BROAD: you cannot picture any concrete answer at all. 'How might we improve healthcare?' invites nothing actionable.",
    "JUST RIGHT: many different answers are imaginable, and you could test each of them. 'How might we make a first screening feel safe rather than exposing?'",
])

h2("Problem Statement, Point of View and How Might We")
table(["Artefact", "What it is", "Example"],
      [["Problem statement", "A neutral description of what is wrong", "Elderly patients miss their screening appointments."],
       ["Point of View (POV)", "[User] needs [need] because [insight]", "Mdm Tan needs to feel safe asking questions because she fears being judged for not understanding."],
       ["How Might We (HMW)", "An invitation to generate solutions", "How might we make a first screening feel safe rather than exposing?"],
       ["The failure mode", "A solution disguised as a problem", "'We need an SMS reminder system.' That is an answer, not a problem."]])
p("The Point of View is where raw research becomes a designable problem. Skip it and your How Might "
  "We question floats free of evidence — it will sound reasonable and lead nowhere.")

h2("The Problem-Assumption Model")
img("problem-assumption.png", "The Problem-Assumption Model (Schneider & O'Reilly) — the bridge from Design Thinking into Lean validation.")
p("The Problem-Assumption Model, created by Jonny Schneider and Barry O'Reilly, is a four-question "
  "worksheet that converts a belief into a test:")
numbered([
    "What's the problem?",
    "How might we solve it?",
    "What assumptions have we made?",
    "How will we test our assumptions?",
])
p("It rests on a three-step learning discipline: define your beliefs and assumptions so that they "
  "can be tested; decide the most important thing to learn; then design the experiments that will "
  "deliver that learning.")
note("Every assumption you cannot test is a risk you have chosen not to see.")

h2("Empathy Maps")
img("empathy-map.png", "The empathy map — four quadrants plus pains and gains.")
p("An empathy map captures four things about a user, and two more at the base:")
dl([("Says", "Direct quotes — what the user says out loud. Observable and quotable."),
    ("Thinks", "What the user believes but does not say out loud. This is inference; label it as such."),
    ("Does", "Observable behaviour, either in general or in response to a specific trigger."),
    ("Feels", "The user's emotional state — for example, 'is confused by the navigation and blames themselves'."),
    ("Pains", "Obstacles and frustrations worth considering — unfamiliarity with technology, a short attention span, fear of a bad result."),
    ("Gains", "What the user hopes to accomplish, and what success looks like in their own terms.")])
note("The most valuable thing on an empathy map is a contradiction between what the user SAYS and "
     "what they DO. That gap is almost always where the real insight lives.")
h3("Common empathy methods")
dl([("Interviews", "Gathering in-depth insight through direct conversation. Ask why repeatedly."),
    ("Observation", "Watching the task performed in its real context, without intervening."),
    ("Questionnaires", "Collecting structured data from a larger audience — breadth, not depth."),
    ("Focus groups", "Facilitating dynamic group discussion to surface shared and contested views."),
    ("Workshop sessions", "Collaborative problem-solving and idea generation with the users themselves.")])

h2("Persona Mapping")
p("User personas are representations of your target customers, built by researching and outlining "
  "your ideal customer's goals, pain points, behaviours and demographic information.")
bullets([
    "Create three to five detailed personas to start — enough to show meaningful variation, few enough for the team to hold in mind.",
    "Ground every claim. Each goal, pain and behaviour should trace back to something a real person said or did.",
    "Include the awkward persona. The user who does not fit the happy path is usually where your design breaks.",
    "Keep layouts consistent between personas and find a common metric to track across them, so they can be compared.",
    "Use icons and visuals to make the data memorable, and present all personas on one page for easy comparison.",
    "Label AI-generated personas as hypotheses. An AI persona is a compression of what the internet says, not evidence about your users.",
    "Retire personas honestly. A persona contradicted by evidence must be changed, not defended.",
])

h2("Ideation Techniques")
dl([("Brainstorming", "Group generation under explicit rules — defer judgement, encourage wild ideas, build on the ideas of others, go for quantity, stay focused, one conversation at a time, be visual."),
    ("Brainwriting", "Silent written generation before any discussion, which prevents the loudest voice in the room from anchoring everybody else."),
    ("Crazy 8s", "Eight rapid variations in eight minutes. The speed is the mechanism: it defeats the internal critic."),
    ("Worst Possible Idea", "Deliberately generate terrible ideas to break the fear of contributing, then invert the best of the worst."),
    ("SCAMPER", "Substitute, Combine, Adapt, Modify, Put to another use, Eliminate, Reverse — a systematic prompt set for variation."),
    ("AI amplification", "Generate manually first, then use GenAI to extend the set. Never the reverse.")])
note("The one rule of ideation: separate the generation of ideas from the evaluation of ideas. "
     "Judging while generating kills the unusual ideas first — and the unusual idea is the one you "
     "came for.")

h2("Prototyping and Testing")
p("A prototype's job is to buy information at the lowest possible price. It is built to think, not "
  "built to ship.")
bullets([
    "Low fidelity invites honesty. A rough paper screen gets real objections; a polished mockup gets polite notes about the font.",
    "Wizard of Oz: fake the mechanism with a human behind the curtain, and learn the behaviour before building anything.",
    "Storyboard the experience as a strip so the team can role-play it and feel where it breaks.",
    "Write the falsification criteria in advance. State what result would prove the concept wrong — otherwise you have a demonstration, not a test.",
    "Use GenAI for the drafting layer: screen copy, user flows, edge cases and interview scripts. You still run the test with real people.",
])

# ================================================================= TOPIC 3
rule()
h1("Topic 3 — Agile Development and AI for Rapid Solution Delivery")
p("Topic weighting: 25%. This topic covers how a validated problem becomes a backlog, a sprint and "
  "a delivered increment — and where generative AI genuinely speeds that up.")

h2("Being Agile versus Doing Agile")
table(["", "Doing Agile (the rituals)", "Being Agile (the mindset)"],
      [["Stand-up", "A status report to the manager", "The team re-plans its own day"],
       ["Backlog", "A queue of requirements handed down", "A living, validated set of bets"],
       ["Sprint review", "A demo performed for stakeholders", "A genuine request for disconfirming feedback"],
       ["Retrospective", "A meeting that produces a tidy list", "The team says the uncomfortable thing and changes something"],
       ["Change request", "A disruption to be resisted", "New information to be welcomed and evaluated"],
       ["Velocity", "A productivity target imposed on the team", "A planning input owned by the team"]])
quote("Practices without principles are a short-lived Band-Aid. And principles without practices are "
      "a fruitless exercise in philosophy.")

h2("Scrum")
img("scrum-loop.png", "The Scrum loop — backlog to sprint to increment, with the retrospective feeding the next cycle.")
p("Scrum is a framework for delivering the highest value in the shortest time through an empirical "
  "process built on transparency, inspection and adaptation. It is suited to complex product "
  "development where requirements will change.")
h3("What Scrum is not")
bullets([
    "It is not undisciplined. Inspection and adaptation, just-in-time planning and solid engineering practices all take real discipline.",
    "It is not a detailed Gantt chart plan.",
    "It is not a silver bullet. Every project differs — stick to the core values and accountabilities when applying it.",
])
h3("The three accountabilities")
dl([("Product Owner", "Owns the vision for the product, creates and maintains the Product Backlog, and is the final decision-maker on prioritisation. Owns the WHAT."),
    ("Scrum Master", "Facilitates the process, builds a self-organising team, removes impediments, protects the team from external disturbance and coaches the practice."),
    ("Developers", "The cross-functional group who build the increment — design, build, test and everything else needed for a potentially shippable result. They own the HOW.")])
p("A team of around seven people, plus or minus two, is the usual guidance. The team is "
  "cross-functional, self-organising and self-managing; it plans its own sprint, swarms on tasks to "
  "minimise idle work, and communicates face-to-face wherever possible.")

h2("Dual-Track Agile")
img("dual-track-agile.png", "Dual-track agile — discovery and delivery running simultaneously in one team.")
p("Dual-track agile was first documented by Desiree Sy in 2007 and popularised by Jeff Patton and "
  "Marty Cagan. It runs two tracks at once inside a single team:")
dl([("The discovery track", "Focuses on fast learning and validation. It asks: should we build this, and what exactly? It produces validated product backlog items."),
    ("The delivery track", "Focuses on predictability and quality. It asks: how do we build this well? It produces releasable increments.")])
quote("The most expensive way to test your idea is to build production quality software.")
p("Discovery has three legitimate outcomes: build forward, kill the idea, or keep learning. If you "
  "are doing discovery right, you will substantially change and kill a lot of ideas — that is the "
  "point, not a failure.")
h3("Four myths about dual-track")
numbered([
    "'Discovery comes first, then delivery.' No — both run continuously and simultaneously. Discovery is a necessary part of product development, practised with the same agile and lean principles.",
    "'All work flows from discovery into development.' No — ideas are frequently abandoned during discovery.",
    "'They are separate teams.' No — a product manager, designer and senior engineer may lead discovery, but they must involve the whole team wherever possible.",
    "'Discovery ends at launch.' No — keep measuring and learning after you ship.",
])
note("The anti-pattern to watch for is what Marty Cagan calls doing little mini-waterfalls inside "
     "your Scrum framework: the product manager writes requirements, hands them to a designer who "
     "produces wireframes, who hands them to the delivery team. Each handoff looks like "
     "collaboration and is actually a queue. The fix is the product manager, designer and lead "
     "engineer working together, side by side, to create and validate backlog items.")

h2("Converting Ideas into Agile Artefacts")
table(["Level", "What it is", "Example"],
      [["Epic", "A large body of work spanning many sprints", "Self-service health screening"],
       ["Feature", "A coherent capability within an epic", "Appointment booking without a call centre"],
       ["Hill", "An outcome statement: as a [who], I want [what], so that [wow]", "As a first-time patient, I can book a screening in under 3 minutes, so that I never need to phone."],
       ["User story", "A unit of value: as a [user], I want [goal], so that [reason]", "As a patient, I want to see the next three available slots so that I can choose quickly."],
       ["Task", "The work needed to deliver the story", "Build the slot-availability API endpoint."]])
p("Hills come from IBM Design Thinking. A Hill states the outcome and the 'wow'; a user story states "
  "the function. Teach and use both — they do different jobs. Three Hills per release is the usual "
  "guidance, and each must be measurable. In one documented IBM session, a manufacturing user's need "
  "for inventory data started as a vague 'in minutes' and was pushed until it became '5 minutes', "
  "justified because the data had to be available before the morning manager meeting. That "
  "specificity is what separates a testable outcome from an aspiration.")

h3("User stories — the three C's")
dl([("Card", "The story is written small, on a card. Its size is a deliberate constraint on scope."),
    ("Conversation", "The card is a promise to have a conversation. The detail lives in that discussion, not in the text."),
    ("Confirmation", "Acceptance criteria define how you will know it is done — testable conditions, not a restatement of the story.")])
note("The most common acceptance-criteria failure, and the one generative AI reproduces most "
     "reliably, is criteria that merely repeat the story in other words. If it cannot fail a test, "
     "it is not an acceptance criterion.")

h2("Release Prioritisation — The Three-Cake Model")
img("cake-releases.png", "Cupcake, Birthday Cake, Wedding Cake — every release is a whole cake at its own scale.")
p("IBM Design Thinking prioritises stories across three releases using an unusually memorable "
  "metaphor. The Cupcake release is complete and satisfying at small scale. The Birthday Cake "
  "release serves more people with broader value. The Wedding Cake release is the full realised "
  "vision.")
p("The point of the metaphor is that a cupcake is a WHOLE cake — not a slice of an unfinished "
  "wedding cake. A partially built product is not a smaller product; it is a broken one, and "
  "shipping it teaches you nothing except that users dislike broken things. Contrast this with the "
  "cost of skipping validation entirely: a builder spends months and thousands of dollars on a "
  "handmade steel bookshelf, only to discover the customer bought an e-reader and no longer reads "
  "paper books.")

h2("Agile Estimation")
dl([("Size, not duration", "Story points measure relative complexity and effort, not hours. Duration is derived later from observed velocity."),
    ("Relative, not absolute", "Humans compare well and predict poorly. 'This is twice that' is far more reliable than 'this takes six hours'."),
    ("A constrained scale", "Values based on the Fibonacci sequence (1, 2, 3, 5, 8, 13, 20, 40, 100) force meaningful distinctions rather than false precision."),
    ("Additive", "Points can be summed across stories in a way that time-based estimates cannot honestly be."),
    ("Velocity from done work", "Only fully 'done, done' stories count. Partially completed stories count for nothing — the last 10% routinely takes 90% of the time, and business value is not achieved until it is done."),
    ("Planning Poker", "Estimate privately, reveal simultaneously, then discuss the outliers. That discussion — not the number — is the real value.")])
p("The actual velocity of the last two iterations is the planned velocity of the next. Sizing must "
  "account for all the cross-functional effort involved: design, code, test, copywriting, "
  "documentation and localisation.")

h2("Ceremonies, Artefacts and Information Radiators")
h3("The daily stand-up")
p("Fifteen minutes maximum. Each person reports what they did since the last stand-up, what they "
  "will do before the next one, and what is blocking them. There is no discussion or debate during "
  "the stand-up — listening only. Problem-solving starts after the meeting ends, and the Scrum "
  "Master then leads the removal of blockers.")
h3("Artefacts")
bullets([
    "Product Backlog — a broad, prioritised list of all required features and wish-list items, with rough estimates of business value and development effort. It is the WHAT that will be built.",
    "Sprint Backlog — the prioritised list of stories selected for a given sprint.",
    "Sprint Burndown Chart — completion of story points over the sprint.",
    "Release Burndown Chart — progress across a whole release.",
])
h3("Information radiators")
p("An information radiator, a term coined by Alistair Cockburn, is a highly visible graphical "
  "display of project status in the team's workspace: the current iteration's stories, work "
  "assignments, test counts, delivered stories and actions from the previous retrospective. It keeps "
  "the team focused and drives transparency across the organisation.")
h3("Retrospective techniques")
bullets(["Post-it brainstorming", "Dot voting", "Timeline", "Team radar", "Keep / Drop / Try"])

h2("GenAI Across the Sprint")
dl([("Backlog drafting", "Generate candidate stories and acceptance criteria from a Hill in seconds — then cut everything that assumes something you never validated."),
    ("Backlog grooming", "Flag duplicate and related items across a large backlog faster than any human scan."),
    ("Test data", "Generate realistic edge-case data sets without exposing real customer records."),
    ("Retrospective synthesis", "Cluster themes across many notes — after the team has spoken, never instead of it."),
    ("Sprint reporting", "Draft the summary so the Scrum Master spends the time on impediments rather than formatting."),
    ("Where it must not run", "Prioritisation, the mid-sprint disruption decision, and the retrospective conversation itself.")])

# ================================================================= TOPIC 4
rule()
h1("Topic 4 — Scaling and Sustaining Innovations with Agile Design Thinking and Generative AI")
p("Topic weighting: 25%. This topic addresses the hardest part: making innovation survive contact "
  "with the wider organisation, and proving it worked.")

h2("The Three Levers That Actually Scale Innovation")
dl([("Purpose, alignment and autonomy", "Be stubborn on the vision, but flexible on the details. Teams need a clear outcome and genuine authority to pursue it. Visualising the whole end-to-end process — from aspirations and hypotheses through design experiments to feedback — on a large product wall lets the whole team play along together."),
    ("Measure things that matter", "Structure your metrics around the future decisions you want to make. Only measure things that indicate progress toward your goal."),
    ("Make decisions based on learning", "Define beliefs and assumptions so they can be tested, decide the most important thing to learn, then design experiments that deliver that learning.")])
quote("If a measurement matters at all, it is because it must have some conceivable effect on "
      "decisions and behaviour.")
p("Hypothesis-Driven Development offers a repeatable format for framing outcomes, beliefs and "
  "metrics, which makes them easy to communicate to others.")

h2("Culture — The Four Principles")
p("Steve Perkins argues that new practices and methodologies only succeed in a lasting way if the "
  "underlying culture supports them. Four principles carry that culture:")
numbered([
    "People over process — assemble people around a shared mission rather than a bureaucratic structure, and prioritise real-time collaboration.",
    "Creativity and innovation — allow personal freedom for self-expression, encourage smart risk-taking, and build exploration time into the workflow.",
    "Iteration, prototyping and adaptation — replace year-long waterfall-style projects with low-fidelity versions tested in real environments.",
    "Multi-disciplinary autonomous teams — self-managing cross-functional teams with genuine local decision-making authority beat hierarchical structures.",
])
h3("The rollout playbook")
bullets([
    "Run design thinking and agile in PARALLEL, not sequentially.",
    "Emphasise principles and practices simultaneously.",
    "Translate concepts into simple, conversational language rather than jargon.",
    "Reinforce continuously so it does not become the flavour of the month.",
    "Give specific behavioural examples of what good looks like.",
    "Celebrate small wins.",
    "Start small: choose low-risk, high-value opportunities and attempt one to three before scaling.",
    "Enable cross-functional collaboration, including physical gatherings with real end-users.",
    "Help agile teams understand the value of the ideation, definition and empathy phases, and allow reframing before development proceeds.",
])
note("Design thinking requires an organisational culture that is open, trusting and encouraging. "
     "Without that foundation the methodology struggles to take root. Teaching design thinking "
     "broadly builds valuable awareness, but it does not replace the need for designers with deep "
     "expertise, and designers need a seat at the table from the outset.")

h2("Engaging Stakeholders and Creating Buy-In")
bullets([
    "Start small and visible — one credible win buys more permission than any presentation.",
    "Bring evidence, not opinion: the user quote, the test result, and the assumption you falsified.",
    "Agree the success metric BEFORE the pilot runs, not after the results are in.",
    "Involve the sceptic early. The person most likely to block it is the most valuable participant in discovery.",
    "Have leaders participate rather than sponsor. A leader who sits in a user interview changes behaviour more than a memo.",
    "State what you will stop doing if the evidence goes against you. Naming the downside earns credibility.",
])

h2("Resource Management in Innovation Projects")
bullets([
    "Fund learning rather than plans — release budget in stages tied to validated learning.",
    "Protect discovery capacity. If discovery is unstaffed, delivery will build whatever is loudest.",
    "Use AI to free expert time, so scarce expertise goes to judgement work rather than synthesis.",
    "Track the cost of being wrong, not just the cost of the test. The relevant number is what building the wrong thing would cost.",
    "Govern AI use: record provenance, check for bias, keep an audit trail for AI-assisted decisions.",
    "Kill weak work deliberately. Stopping an initiative releases your scarcest resource — your best people.",
])

h2("Sensemaking")
p("Sensemaking is the discipline of turning ambiguous, conflicting and partial signals into a shared "
  "narrative that an organisation can act on. Real innovation decisions are never made on complete "
  "evidence.")
bullets([
    "Signals arrive incomplete and contested — waiting for certainty is itself a decision.",
    "The output is a shared narrative the organisation can act on together, not a data dump.",
    "The evidence that contradicts your narrative is the most valuable input in the room.",
    "Generative AI clusters signals well, and will confidently manufacture a coherent narrative even when the underlying evidence is thin. Fluency is not truth.",
])

h2("Systems Thinking")
img("feedback-loops.png", "Reinforcing and balancing feedback loops.")
p("Systems thinking sees the whole rather than the parts. Five ideas carry most of its practical "
  "value:")
dl([("Interconnectedness", "Shift from linear cause-and-effect to circular thinking. A system is a set of components working together to achieve an objective; everything is part of a wider ecosystem."),
    ("Synthesis", "Combining two or more connections to create a new idea or concept. It is the opposite of analysis, which dissects complexity into parts and loses the behaviour."),
    ("Emergence", "The natural outcome of things coming together. Emergent behaviour is non-linear and self-organised; it cannot be predicted from the parts alone."),
    ("Feedback loops", "With connection come loops and flows. Reinforcing loops amplify a change; balancing loops stabilise it. Both are observable and can be altered."),
    ("Causality", "Understanding how a change in one variable actually propagates to an outcome in a dynamic environment.")])
p("Systems mapping draws the elements and their interconnections to reveal insights, guide policy "
  "decisions and identify the leverage point where an intervention will actually work.")
note("Systemic failure is typically a reinforcing loop with no effective balancing counterpart. No "
     "individual decision in the chain has to be irrational for the outcome to be catastrophic — "
     "which is exactly why individual good judgement is not a sufficient safety mechanism.")
table(["", "Design Thinking", "Systems Thinking"],
      [["Starting point", "A human need", "A pattern of behaviour over time"],
       ["Unit of attention", "The user and their experience", "The structure that produces the outcome"],
       ["Asks", "What does this person need?", "Why does this system keep producing this result?"],
       ["Typical output", "A validated solution concept", "A leverage point and an intervention"],
       ["Blind spot", "Can optimise for one user while harming the system", "Can map elegantly and never ship anything"],
       ["Used together", "Design the intervention with empathy", "Place it where the system will actually respond"]])

h2("Metrics and KPIs for Innovation")
img("metrics-quadrant.png", "The innovation metric test — does it change a decision, and is it timely?")
p("The decisive test for any innovation metric is Hubbard's: if a movement in this number would "
  "change no decision that anyone would make, the number is decoration. Vanity metrics — ideas "
  "submitted, workshops run, people trained — always rise and prove nothing, which is precisely why "
  "they are popular and why they do not survive a budget review.")
h3("A balanced scorecard")
dl([("Customer value", "Time saved, effort removed, task success rate, or a genuine satisfaction signal."),
    ("Business value", "Revenue, cost avoided, retention or cycle time, expressed in money where that is honest."),
    ("Societal value", "Access, inclusion, safety or environmental effect — the value not captured in the price."),
    ("A leading indicator", "Something that moves early enough to still change a decision this quarter.")])
h3("Three separate validations")
img("three-validations.png", "Problem, solution and demand validation are three different tests.")
p("Do not conflate a good prototype test result with a strong affinity for the problem, or with "
  "customer demand for the solution. These are separate concerns requiring different learning "
  "approaches:")
dl([("Problem validation", "Does this problem genuinely exist, and does it matter to enough people?"),
    ("Solution validation", "Does this solution actually solve the problem for those people?"),
    ("Demand validation", "Will they adopt it, switch to it, or pay for it?")])
note("Teams routinely present solution-validation evidence as though it proved demand. That is how "
     "confidently-built products launch to silence.")

# ================================================================= ACTIVITIES
rule()
h1("Hands-On Activities")
p("The twelve activities below are the practical core of this course. Each is built on a documented "
  "real-world case, and each follows the same structure: the case, your scenario, the step-by-step "
  "instructions, the discussion questions and the debrief.")
p("Work in groups of three to five. Each activity also exists as a standalone PDF brief in the "
  "activities folder, so you can work from a printed copy at your table.")

TOPIC_BY_NUM = {t["num"]: t for t in C.TOPICS}
for a in ACT:
    tp = TOPIC_BY_NUM[a["topic"]]
    h2(f"Activity {a['num']} — {a['title']}")
    p(f"Topic {tp['code']}: {tp['title']}   ·   Duration: {a.get('duration', 45)} minutes   ·   "
      f"Tools: {a['services']}")
    h3("Learning outcome addressed")
    p(a["objective"])
    h3("The real case")
    p(a["case"])
    h3("Your scenario")
    p(a["scenario"])
    h3("What you will do")
    p(a["desc"])
    h3("What you will produce")
    p(a["build"])
    h3("Step-by-step instructions")
    steps(a["steps"])
    h3("Discussion questions")
    numbered(a["discussion"])
    h3("How you know you are done")
    p(a["test"])
    h3("Debrief — what the room should conclude")
    p(a["debrief"])
    rule()

# ================================================================= ASSESSMENT
h1("Preparing for the Assessment")
p("The final assessment has two instruments, both open book and both conducted at the end of Day 2.")
bullets([C.ASSESSMENT["written"], C.ASSESSMENT["practical"],
         "You may use this Learner Guide, the course slides, your activity briefs and any approved material.",
         "You are assessed as Competent (C) or Not Yet Competent (NYC) on each instrument.",
         C.ASSESSMENT["note"]])
h3("What to revise")
bullets([
    "The three mindsets and the question each answers — Design Thinking, Lean and Agile.",
    "The Double Diamond, and which question each diamond answers.",
    "The five stages of design thinking and what happens in each.",
    "The difference between a problem statement, a Point of View and a How Might We question, and the altitude test.",
    "The Problem-Assumption Model's four questions.",
    "Empathy map quadrants, and why the say/do contradiction matters.",
    "Dual-track agile: the two tracks, what each asks, and the mini-waterfall anti-pattern.",
    "The decomposition ladder from Epic to Task, and the difference between a Hill and a user story.",
    "Story points versus duration, and how velocity is counted.",
    "The three-cake release model and what makes a Cupcake a whole cake.",
    "Reinforcing versus balancing feedback loops.",
    "The metric decision test, and the three separate validations (problem, solution, demand).",
    "Where generative AI genuinely helps in the loop, and where it must not be trusted.",
])
h3("Assessment day flow")
numbered([
    "Complete the mandatory TRAQOM course feedback survey on the LMS.",
    "Take the Assessment digital attendance by scanning the SSG QR code.",
    "Sit the Written Assessment (SAQ), then the Case Study — one hour each, open book.",
    "Submit your completed answers on the LMS at https://lms-tms.tertiaryinfotech.com.",
    "Sign the Assessment Summary Record.",
])
note("All five steps are mandatory for WSQ funding. Missing the digital attendance or the TRAQOM "
     "survey can invalidate your funding claim.")

# ================================================================= GLOSSARY
h1("Glossary")
dl([
 ("Acceptance criteria", "Testable conditions that define when a user story is done. Not a restatement of the story."),
 ("Agile", "A mindset and set of frameworks for adapting to changing conditions while delivering incrementally."),
 ("Balancing loop", "A feedback loop that stabilises a system by counteracting change."),
 ("Brainwriting", "Silent, written idea generation used before discussion to prevent anchoring."),
 ("Crazy 8s", "An ideation technique producing eight rapid variations in eight minutes."),
 ("Cupcake / Birthday Cake / Wedding Cake", "IBM's three-release prioritisation model; each release is a complete product at its own scale."),
 ("Design Sprint", "A time-boxed cross-functional activity that takes a big problem to a clear, tested direction which can then feed an agile backlog."),
 ("Design Thinking", "A human-centred, solution-based approach to exploring and solving ill-defined problems."),
 ("Discovery track", "The continuous validation track in dual-track agile; asks whether we should build something and what exactly."),
 ("Delivery track", "The build track in dual-track agile; asks how to build the validated thing well."),
 ("Double Diamond", "A model of innovation as two linked diamonds — the problem space and the solution space — each with a diverge and converge phase."),
 ("Dual-track agile", "Running discovery and delivery simultaneously within a single product team."),
 ("Emergence", "System behaviour that arises from interaction between parts and cannot be predicted from the parts alone."),
 ("Empathy map", "A four-quadrant tool capturing what a user says, thinks, does and feels, plus their pains and gains."),
 ("Epic", "A large body of work spanning many sprints, decomposed into features and stories."),
 ("Hill", "An IBM Design Thinking outcome statement in the form: as a [who], I want [what], so that [wow]."),
 ("How Might We (HMW)", "A question format that reframes a problem as an invitation to generate solutions."),
 ("Hypothesis-Driven Development", "A format for framing outcomes, beliefs and the metrics that would confirm or refute them."),
 ("Information radiator", "A highly visible display of project status in the team's workspace (Alistair Cockburn)."),
 ("Lean Startup", "An approach that validates business ideas cheaply through minimum viable products and customer feedback."),
 ("Minimum Viable Product (MVP)", "The smallest product that generates real learning about customer value."),
 ("Persona", "A research-grounded representation of a target user's goals, pains, behaviours and context."),
 ("Point of View (POV)", "A designable problem statement in the form: [user] needs [need] because [insight]."),
 ("Product Backlog", "The prioritised list of everything that might be built, owned by the Product Owner."),
 ("Product Owner", "The Scrum accountability that owns the product vision, the backlog and prioritisation."),
 ("Prototype", "A cheap, low-fidelity artefact built to think with and to test an idea, not to ship."),
 ("Reinforcing loop", "A feedback loop that amplifies change, driving a system further in one direction."),
 ("Retrospective", "A recurring team meeting to inspect how the team works and commit to changes."),
 ("SCAMPER", "An ideation prompt set: Substitute, Combine, Adapt, Modify, Put to another use, Eliminate, Reverse."),
 ("Scrum", "An agile framework delivering value through short sprints with defined accountabilities, events and artefacts."),
 ("Scrum Master", "The Scrum accountability that enables the process, removes impediments and coaches the team."),
 ("Sensemaking", "Turning ambiguous, conflicting signals into a shared narrative an organisation can act on."),
 ("Sprint", "A fixed-length iteration, typically one to four weeks, producing a potentially shippable increment."),
 ("Story point", "A relative measure of the size and complexity of a user story, not its duration."),
 ("Systems thinking", "Seeing a problem as a whole system of interconnections, feedback loops and emergent behaviour."),
 ("Three C's", "Card, Conversation, Confirmation — the components of a well-formed user story."),
 ("User story", "A small unit of value in the form: as a [user], I want [goal], so that [reason]."),
 ("Vanity metric", "A number that reliably rises but changes no decision."),
 ("Velocity", "The amount of fully completed work a team delivers per sprint, used as a planning input."),
 ("Wicked problem", "A problem with interconnected layers, no fixed endpoint and no single correct solution (Rittel & Webber, 1973)."),
 ("Wizard of Oz prototype", "A prototype where a human secretly performs the function, to test behaviour before building."),
])

h1("References and Further Reading")
bullets([
 "Schneider, J. — Understanding how Design Thinking, Lean and Agile Work Together. Thoughtworks.",
 "Adaptovate — The Relationship Between Design Thinking and Agile.",
 "BMC Blogs — Design Thinking vs Lean vs Agile.",
 "Burba, D. (2016) — Agile by Design: Integrating Design Thinking and Agile Approaches Helps Organizations Find and Build the Right Customer-Focused Solution. PM Network, 30(10), 58–63.",
 "Vukosav, D. (2019) — Design Thinking to Improve Your Agile Process. PMI Global Congress EMEA, Dublin.",
 "UX Magazine — Agile and Design Thinking: How Can They Go Well Together.",
 "Perkins, S. — Agile & Design Thinking: Competing or Completing? The Design Gym.",
 "PremierAgile — Design Thinking vs Agile.",
 "Startup Frontier (Medium) — How to Combine Design Thinking and Agile in Practice.",
 "Rittel, H. & Webber, M. (1973) — Dilemmas in a General Theory of Planning (wicked problems).",
 "Sy, D. (2007) — Adapting Usability Investigations for Agile User-Centered Design.",
 "Patton, J. — Dual Track Development is not Duel Track.",
 "Cagan, M. — Dual-Track Agile. Silicon Valley Product Group.",
 "Liedtka, J. — Seven-year study of design thinking in practice, University of Virginia Darden School.",
 "Hubbard, D. W. — How to Measure Anything: Finding the Value of Intangibles in Business.",
])


# ---------------- render Markdown ----------------
def _anchor(txt):
    return "".join(ch.lower() if ch.isalnum() else ("-" if ch in " -" else "") for ch in txt)


def render_md():
    out = [f"# {C.TITLE} — Learner Guide", ""]
    out.append(f"**WSQ Course Code:** {C.COURSE_CODE}  |  **Conducted by:** {C.ORG} "
               f"({C.UEN.replace('UEN: ', 'UEN ')})  |  **Version {C.VERSION} · {C.VERSION_DATE}**")
    out.append("")
    out.append("## Contents"); out.append("")
    for kind, *rest in B:
        if kind == "h1":
            out.append(f"- [{rest[0]}](#{_anchor(rest[0])})")
        elif kind == "h2":
            out.append(f"  - [{rest[0]}](#{_anchor(rest[0])})")
    out.append("")
    for kind, *rest in B:
        if kind == "h1":
            out += ["", f"## {rest[0]}", ""]
        elif kind == "h2":
            out += ["", f"### {rest[0]}", ""]
        elif kind == "h3":
            out += [f"**{rest[0]}**", ""]
        elif kind == "p":
            out += [rest[0], ""]
        elif kind == "bullets":
            out += [f"- {x}" for x in rest[0]] + [""]
        elif kind == "numbered":
            out += [f"{i}. {x}" for i, x in enumerate(rest[0], 1)] + [""]
        elif kind == "steps":
            for i, (instr, cmd) in enumerate(rest[0], 1):
                out.append(f"{i}. {instr}")
                if cmd:
                    out += ["", "   ```", f"   {cmd}", "   ```", ""]
            out.append("")
        elif kind == "note":
            out += [f"> **Note:** {rest[0]}", ""]
        elif kind == "quote":
            out += [f"> {rest[0]}", ""]
        elif kind == "table":
            hdrs, rows = rest[0], rest[1]
            out.append("| " + " | ".join(hdrs) + " |")
            out.append("|" + "|".join(["---"] * len(hdrs)) + "|")
            for r in rows:
                out.append("| " + " | ".join(str(c) for c in r) + " |")
            out.append("")
        elif kind == "img":
            name, cap = rest[0], rest[1]
            rel = os.path.join(".claude/skills/courseware-build/assets/img", name)
            out += [f"![{cap}]({rel})", ""]
            if cap:
                out += [f"*{cap}*", ""]
        elif kind == "rule":
            out += ["---", ""]
        elif kind == "dl":
            for term, defn in rest[0]:
                out.append(f"- **{term}** — {defn}")
            out.append("")
    return "\n".join(out)


MD_OUT = os.path.join(REPO, f"LG-{C.SHORT_TITLE}.md")
with open(MD_OUT, "w") as f:
    f.write(render_md())
print("Saved", MD_OUT)

# ---------------- render DOCX ----------------
BRAND = RGBColor(0x1F, 0x6F, 0xEB); DARK = RGBColor(0x11, 0x18, 0x27); GREY = RGBColor(0x55, 0x5B, 0x66)
doc = Document()
normal = doc.styles["Normal"]; normal.font.name = "Arial"; normal.font.size = Pt(11)
prodoc.style_headings(doc)
prodoc.add_cover_page(doc, "LEARNER GUIDE", C.TITLE, C.VERSION.lstrip("v"),
                      org_logo=os.path.join(ASSETS, "tertiary-infotech-logo.png"),
                      course_logo=None, course_code=C.COURSE_CODE)
prodoc.add_version_control(doc, [
 ("3.0", "1 January 2026",
  "Legacy learner material accompanying the v3 master trainer deck.",
  C.TRAINER),
 ("4.0", C.VERSION_DATE,
  "Major revision. Teaching content substantially expanded from Thoughtworks, Adaptovate, BMC, "
  "PMI, UX Magazine, The Design Gym, PremierAgile and Startup Frontier — adding the three-mindset "
  "model, the Double Diamond, wicked problems, the Problem-Assumption Model, dual-track agile, "
  "Hills, the three-cake release model, systems thinking and the innovation metric test. Twelve "
  "real-world case-study activities with full step-by-step instructions, discussion questions and "
  "debriefs. Design Thinking Studio and Padlet Classroom Board adopted as the class ed-tools.",
  C.TRAINER),
])
prodoc.add_toc(doc)


def _add_table(headers, rows):
    t = doc.add_table(rows=0, cols=len(headers)); t.style = "Table Grid"
    hdr = t.add_row().cells
    for i, htext in enumerate(headers):
        hdr[i].text = ""
        r = hdr[i].paragraphs[0].add_run(str(htext))
        r.bold = True; r.font.size = Pt(9.5); r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        prodoc._shade_cell(hdr[i], "1F6FEB")
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            r = cells[i].paragraphs[0].add_run(str(val)); r.font.size = Pt(9.5)
            if i == 0:
                r.bold = True
                prodoc._shade_cell(cells[i], "E8F0FE")
    doc.add_paragraph("")


for kind, *rest in B:
    if kind == "h1":
        doc.add_heading(rest[0], level=1)
    elif kind == "h2":
        doc.add_heading(rest[0], level=2)
    elif kind == "h3":
        para = doc.add_paragraph()
        r = para.add_run(rest[0]); r.bold = True; r.font.size = Pt(11); r.font.color.rgb = BRAND
    elif kind == "p":
        doc.add_paragraph(rest[0])
    elif kind == "bullets":
        for x in rest[0]:
            doc.add_paragraph(x, style="List Bullet")
    elif kind == "numbered":
        for i, x in enumerate(rest[0], 1):
            para = doc.add_paragraph(); para.paragraph_format.left_indent = Inches(0.3)
            r = para.add_run(f"{i}. "); r.bold = True
            para.add_run(x)
    elif kind == "steps":
        for i, (instr, cmd) in enumerate(rest[0], 1):
            para = doc.add_paragraph(); para.paragraph_format.left_indent = Inches(0.3)
            r = para.add_run(f"{i}. "); r.bold = True
            para.add_run(instr)
            if cmd:
                cp = doc.add_paragraph()
                r = cp.add_run(cmd); r.font.name = "Consolas"; r.font.size = Pt(9.5)
                r.font.color.rgb = RGBColor(0x0B, 0x30, 0x60)
    elif kind == "note":
        para = doc.add_paragraph()
        r = para.add_run("Note: "); r.bold = True; r.font.color.rgb = BRAND
        para.add_run(rest[0]).font.size = Pt(10)
    elif kind == "quote":
        para = doc.add_paragraph()
        r = para.add_run("“" + rest[0] + "”")
        r.italic = True; r.font.size = Pt(10.5); r.font.color.rgb = GREY
    elif kind == "table":
        _add_table(rest[0], rest[1])
    elif kind == "img":
        path = os.path.join(IMGDIR, rest[0])
        if os.path.exists(path):
            doc.add_picture(path, width=Inches(6.1))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            if rest[1]:
                cap = doc.add_paragraph()
                cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                r = cap.add_run(rest[1]); r.italic = True; r.font.size = Pt(9.5)
                r.font.color.rgb = GREY
    elif kind == "rule":
        doc.add_paragraph("")
    elif kind == "dl":
        for term, defn in rest[0]:
            para = doc.add_paragraph(style="List Bullet")
            r = para.add_run(term + " — "); r.bold = True
            para.add_run(defn)

prodoc.add_page_numbers(doc)
prodoc.enable_update_fields(doc)
DOCX_OUT = os.path.join(REPO, "courseware", f"LG-{C.SHORT_TITLE}.docx")
doc.save(DOCX_OUT)
print("Saved", DOCX_OUT)
