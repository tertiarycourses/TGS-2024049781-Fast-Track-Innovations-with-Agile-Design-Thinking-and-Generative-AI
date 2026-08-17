#!/usr/bin/env python3
"""
Build the WSQ assessment set for TGS-2024049781 — four DOCX into assessment/:

    WA (SAQ) - <Title> - v4.docx              6 open-ended knowledge questions (K1–K5)
    Answer to WA (SAQ) - <Title> - v4.docx    model answers, citing the slides/LG
    CS Assessment - <Title> - v4.docx         1 scenario + 4 tasks (A1–A7)
    Answer to CS Assessment - <Title> - v4.docx  marking guide, citing the activities

The structure MIRRORS the original papers pulled from the TMS:
  * Written Assessment (SAQ) — 6 questions, 0.5 hr, K-codes printed on each question.
  * Case Study Assessment    — 1 scenario + 4 tasks, 1 hr, A-codes printed on each task.
Only the content is rewritten, from THIS course's slides, Learner Guide and activities.

House format: WSQ cover page (cover only — assessments carry no version-control
record), then Trainee Information + Instructions + Grading on page 2, then the
questions/scenario. Arial 11. All questions OPEN-ENDED — zero multiple choice.

Run:  python3 build_assessment.py
"""
import os, sys
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH as AL

HERE = os.path.dirname(os.path.abspath(__file__)); sys.path.insert(0, HERE)
import course_data as C
import prodoc


def _find_repo(start):
    env = os.environ.get("COURSE_REPO")
    if env and os.path.isdir(env):
        return env
    d = start
    for _ in range(8):
        d = os.path.dirname(d)
        if os.path.isdir(os.path.join(d, "courseware")) and os.path.isdir(os.path.join(d, "labs")):
            return d
    return os.path.dirname(os.path.dirname(HERE))


REPO = _find_repo(HERE)
OUT = os.path.join(REPO, "assessment")
ASSETS = os.path.join(os.path.dirname(HERE), "assets")
ORG_LOGO = os.path.join(ASSETS, "tertiary-infotech-logo.png")
VER = "v4"
LMS_URL = "https://lms-tms.tertiaryinfotech.com/"

DARK = RGBColor(0x16, 0x1B, 0x26); BRAND = RGBColor(0x1F, 0x6F, 0xEB); GREY = RGBColor(0x55, 0x5B, 0x66)


# ================================================================ WA — 6 questions (K1–K5)
# (question, K-code, [model answer points], source)
WRITTEN = [
    ("Describe two recent trends in design thinking that have influenced innovation practices. "
     "How can generative AI tools such as ChatGPT or an image generator be used to support "
     "ideation within these trends?",
     "K1",
     ["Trend 1 — Continuous discovery: discovery is no longer a phase that runs before delivery; it runs "
      "permanently alongside it, so teams validate ideas while other work is being built.",
      "Trend 2 — AI-assisted discovery: research synthesis that previously took a team weeks now takes "
      "hours, which moves the organisational bottleneck from analysis to judgement.",
      "Other acceptable trends: cross-functional discovery trios (product manager, designer, engineer); "
      "evidence over opinion (a named assumption plus the test that would disprove it); responsible-AI "
      "governance built into the process; the shift from measuring output to measuring outcome.",
      "GenAI in ideation — it multiplies idea volume and range on demand, generates deliberately "
      "unconventional provocations, produces analogies from other industries, and drafts concept copy or "
      "images so an idea becomes discussable in minutes rather than days.",
      "The essential qualification: AI amplifies good design thinking discipline rather than substituting "
      "for it. Generate manually FIRST and then use AI to extend — running AI first anchors the room to "
      "its framing and measurably narrows the range of ideas produced.",
      "A language model generates the statistically likely, not the locally true, so AI output is a "
      "hypothesis to validate, never evidence in itself."],
     "Topic 1 — Latest Trends in Innovation; Where Generative AI Accelerates the Loop"),

    ("Explain the role of innovation management in organisational transformation. Illustrate how "
     "generative AI can streamline the development of an innovation pipeline through idea screening "
     "and clustering.",
     "K2",
     ["Innovation management is the discipline of running innovation as a repeatable operating system "
      "rather than as isolated bright ideas — deciding what gets explored, what gets funded, what gets "
      "killed, and what gets measured.",
      "In transformation it supplies the mechanism that lets an organisation change itself: Netflix "
      "repeatedly ran the innovation loop against its own profitable businesses (removing late fees, then "
      "cannibalising DVD-by-mail with streaming, then licensed content with originals), while Blockbuster "
      "optimised the business it already had and failed.",
      "It requires three levers: purpose, alignment and autonomy; measuring things that matter; and making "
      "decisions based on learning rather than on seniority.",
      "GenAI in idea screening — it can rapidly summarise a large volume of submitted ideas, tag them "
      "against strategic themes, flag duplicates and near-duplicates, and surface the small number that "
      "warrant human review.",
      "GenAI in clustering — it groups a mass of qualitative submissions into coherent themes, which "
      "directly attacks the documented design thinking bottleneck (Liedtka's research found that making "
      "sense of the resulting mass of qualitative data is the significant challenge).",
      "The limitation: screening criteria encode strategy, so a model must not make the go/kill decision. "
      "It prepares the decision; a human owns it and remains accountable for it."],
     "Topic 1 — The Three Mindsets; Topic 4 — Three Levers That Scale Innovation; Activity 2"),

    ("Identify three drivers of organisational growth and explain how generative AI can enhance "
     "stakeholder analysis during the early stages of design thinking.",
     "K3",
     ["Driver 1 — Customer value created: time saved, effort removed, or a task made succeed that "
      "previously failed. DBS measured customer hours saved and eliminated tens of millions of them.",
      "Driver 2 — Speed of validated learning: how quickly the organisation can test a belief and act on "
      "the result, which determines how many bets it can afford to make.",
      "Driver 3 — Organisational capability and culture: autonomy backed by real decision rights, "
      "cross-functional teams, and leaders who participate rather than merely sponsor.",
      "Other acceptable drivers: new market access, retention and lifetime value, cost structure "
      "improvement, and the ability to self-disrupt before a competitor does.",
      "GenAI in stakeholder analysis — it can rapidly synthesise interview transcripts, survey free text, "
      "support tickets and social sentiment into candidate themes; draft an initial stakeholder map with "
      "likely interests, influence and objections; generate interview guides; and produce first-draft "
      "personas to be tested.",
      "The critical caution: an AI-generated persona is a compression of what the internet says about a "
      "group, not evidence about YOUR stakeholders. In the gig-economy example the AI personas were "
      "plausible, internally consistent and wrong — real interviews showed the dominant anxiety was "
      "within-month income smoothing, not saving for a flat. Validate before converging."],
     "Topic 4 — Metrics and Drivers; Topic 1 — GenAI in the Loop; Activities 3 and 10"),

    ("Describe a scenario in which collecting stakeholder feedback via AI tools can improve "
     "understanding of end-user pain points. What are the potential limitations?",
     "K3",
     ["Scenario — a polyclinic group finds elderly patients abandoning health screening appointments. The "
      "team collects hundreds of free-text survey comments, call-centre notes and frontline staff "
      "observations, and uses a generative AI tool to cluster them into recurring themes.",
      "The AI surfaces that comments cluster not around forgetting the appointment (the assumed problem) "
      "but around fear of a bad result, uncertainty about cost, and having nobody to accompany them — "
      "which redirects the design away from yet another reminder system.",
      "The benefit: synthesis across a volume of qualitative evidence that a small team could not read "
      "and cross-reference in the available time, and a first pass that is free of the team's own "
      "confirmation bias about what it expects to find.",
      "Limitation 1 — Fluency is not truth. The model produces confident, well-written output whether or "
      "not the underlying evidence supports it, and a wrong-but-confident synthesis is harder to "
      "challenge than an obviously poor one.",
      "Limitation 2 — It only sees what was captured. Users who never responded, never complained, or "
      "never reached the service at all are invisible in the data, and the AI cannot know they are missing.",
      "Limitation 3 — It cannot observe. It could not notice a 74-year-old hesitating over entering her "
      "NRIC in a public corridor — the insight that redirected the self-service kiosk design.",
      "Limitation 4 — Privacy, confidentiality and bias. Personal or client data must not be pasted into "
      "public AI tools, and model bias can under-weight minority experiences.",
      "Mitigation: label every AI-derived insight as unvalidated, and confirm it against a small number "
      "of real interviews or observations before it drives a decision."],
     "Topic 2 — Empathy Maps and Persona Mapping with GenAI; Activities 3 and 6"),

    ("Choose one agile or design project management tool and describe how generative AI features "
     "within it could support project tracking and team collaboration.",
     "K5",
     ["Name a specific tool — for example a shared design thinking workspace, a Kanban or sprint board, "
      "or a collaborative whiteboard. The Design Thinking Studio and the Padlet Classroom Board used in "
      "this course are acceptable answers.",
      "Backlog drafting — AI converts an outcome statement (a Hill: as a [who], I want [what], so that "
      "[wow]) into candidate user stories with draft acceptance criteria in seconds.",
      "Backlog grooming — it flags duplicate and closely related items across a large backlog faster than "
      "any manual scan, and can suggest a merge.",
      "Progress tracking and reporting — it drafts the sprint summary and status report so the Scrum "
      "Master spends the time on removing impediments rather than on formatting.",
      "Retrospective synthesis — it clusters many Keep / Drop / Try notes into themes after the team has "
      "spoken, making patterns across several sprints visible.",
      "Collaboration — it summarises long comment threads, drafts meeting notes, and translates content "
      "for a multilingual team.",
      "Where it must NOT run: prioritisation, the mid-sprint disruption decision, and the retrospective "
      "conversation itself — the value of a retrospective is the team saying uncomfortable things to each "
      "other, not a tidy summary. AI drafts; humans decide.",
      "Watch for the two characteristic AI failures: well-formed stories describing features nobody "
      "validated, and acceptance criteria that restate the story instead of defining a testable condition."],
     "Topic 3 — GenAI Across the Sprint; Converting Ideas into Agile Artefacts; Activities 8 and 9"),

    ("Discuss how the principles of resource management can be applied in agile environments. Provide "
     "an example of using generative AI to allocate resources more effectively in an innovation project.",
     "K4",
     ["Fund learning rather than plans — release budget in stages tied to validated learning, so money "
      "follows evidence instead of a forecast written before anything was known.",
      "Protect discovery capacity — if the discovery track is unstaffed, the delivery track will build "
      "whatever is loudest. Dual-track agile requires both tracks to be resourced within one team.",
      "Prioritise by relative size, not false precision — story points estimate size and complexity, and "
      "velocity is counted only on fully 'done, done' work, which makes capacity planning honest.",
      "Track the cost of being wrong, not just the cost of the test — the relevant comparison is a "
      "two-afternoon paper prototype against a quarter of engineering spent on the wrong thing. The most "
      "expensive way to test an idea is to build production-quality software.",
      "Kill work deliberately — stopping a weak initiative releases the scarcest resource an organisation "
      "has, which is the time of its best people.",
      "Sequence releases so each is genuinely usable — the Cupcake release is a whole cake at small "
      "scale, not a slice of an unfinished wedding cake, so resources committed always return value.",
      "GenAI example — in an innovation project the team uses AI to synthesise research, draft user "
      "stories and generate test data, which removes roughly a week of specialist analyst time; that "
      "capacity is redeployed onto the six real user interviews only a human can conduct. AI absorbs the "
      "volume work so scarce expertise goes to judgement work.",
      "Governance: record the provenance of AI-assisted decisions, check outputs for bias, and keep an "
      "audit trail — responsible AI is part of resource management, not a separate compliance exercise."],
     "Topic 3 — Agile Estimation; Topic 4 — Resource Management in Innovation Projects; Activity 8"),
]

# ================================================================ CS — scenario + 4 tasks (A1–A7)
CS_TITLE = "\"Kampung Connect\" Community Innovation Programme at Sunrise Health Group"

CS_SCENARIO = [
    "Sunrise Health Group operates eleven community clinics across Singapore. Screening uptake among "
    "residents aged 60 and above has been flat for two years despite three separate revamps of the SMS "
    "reminder system, and the board has now approved a six-month innovation programme called "
    "\"Kampung Connect\" to change that.",

    "You are Priya Raman, the newly appointed Innovation Lead. You report to the Chief Operating Officer "
    "and work across four groups who do not normally work together: Clinic Operations (who own the "
    "appointment system and are protective of throughput), the Data Analytics team (who hold the app and "
    "attendance data), Community Outreach (who run events at the void decks and know residents by name), "
    "and Finance (who control the programme budget and want a business case).",

    "The prevailing view in the organisation is that residents forget their appointments. The operations "
    "team has proposed building an improved reminder application with a confirmation button, and has "
    "asked for two-thirds of the programme budget to do it. Nothing in the current evidence base explains "
    "why three previous reminder revamps produced no measurable change.",

    "Early signals point elsewhere. Outreach staff report that residents often say they will attend and "
    "then do not. Free-text survey comments mention cost uncertainty and a reluctance to attend alone. "
    "One nurse has observed that residents who arrive with a family member almost always complete the "
    "screening, while those who arrive alone frequently leave before it starts.",

    "You have access to generative AI tools for research synthesis, ideation and drafting, and to two "
    "collaborative workspaces used across the organisation. You have six months, a cross-functional team "
    "of seven people, and a board that will decide on continued funding at the end of the programme.",
]

CS_TASKS = [
    ("Task 1: Apply and Advocate Design Thinking Across the Organisation",
     "A1 · A6",
     "Mapped to A1 — Integrate design thinking methodologies into processes to drive innovation across "
     "the organisation, and A6 — Cultivate design thinking as a viable tool and methodology to foster "
     "new innovations.",
     ["Design a project plan showing how design thinking will guide the discovery, ideation, prototyping "
      "and refinement of the Kampung Connect programme. Name the stages and what each will produce.",
      "Explain how you would respond to the operations team's proposal to spend two-thirds of the budget "
      "on a reminder application, and justify your position with reference to the evidence available.",
      "Propose a communication or engagement strategy that promotes design thinking as a credible method "
      "across Clinic Operations, Data Analytics, Community Outreach and Finance.",
      "State specifically how generative AI will be used to support creative ideation and problem "
      "solving, and where you will not permit it to be used."],
     ["Project plan structured on the design thinking loop, with named outputs per stage: Empathise "
      "(interviews and observation at clinics and void decks), Define (empathy map, Point of View "
      "statement, How Might We question), Ideate (structured divergence then dot-voted convergence), "
      "Prototype (low-fidelity, testable within days), Test (with real residents, against pre-stated "
      "falsification criteria).",
      "Explicit use of the Double Diamond: the first diamond establishes whether the problem is actually "
      "forgetting; only then does the second diamond explore solutions. Credit answers that state the "
      "first diamond is the one teams skip under pressure.",
      "Response to the reminder-app proposal: the proposal is a solution presented as a problem "
      "statement. Three reminder revamps producing no change is strong evidence that forgetting is not "
      "the binding constraint. The professional move is not refusal but reframing — fund a short, cheap "
      "discovery phase before committing two-thirds of the budget, and make the release of that budget "
      "conditional on the evidence.",
      "Reference to the GE Adventure Series reframe (from 'make the scanner quieter' to 'make a "
      "frightened child feel brave') or to Airbnb (analytics would never have surfaced the photography "
      "insight) as precedent for reframing before building.",
      "Engagement strategy — start small with one clinic and a low-risk, high-value pilot; have leaders "
      "participate in user interviews rather than merely sponsor the programme; involve the sceptic from "
      "Clinic Operations early because they are the most valuable participant in discovery; translate "
      "the method into plain language rather than jargon; and celebrate small wins visibly.",
      "GenAI use: synthesising free-text survey comments and outreach notes into candidate themes, "
      "drafting interview guides, generating idea volume after the team has diverged manually, and "
      "drafting prototype copy and screens.",
      "GenAI limits: it must not generate the personas that stand in for real residents, must not make "
      "the go/kill decision, and must not run the stakeholder conversations. Every AI-derived insight is "
      "labelled unvalidated until confirmed against real evidence.",
      "Credit any answer noting that generating manually first and using AI to extend avoids anchoring "
      "the room to the model's framing."]),

    ("Task 2: Synthesise Insights to Uncover Resident Needs",
     "A3 · A5",
     "Mapped to A3 — Synthesise information from different sources and stakeholders in order to fully "
     "understand the needs of end users, and A5 — Engage stakeholders during the design thinking process "
     "to uncover the motivations behind their actions and behaviours.",
     ["Develop a resident insight brief that synthesises the quantitative data held by Data Analytics "
      "with the qualitative observations from Community Outreach and clinic staff.",
      "Identify at least three resident pain points that are NOT 'they forget', and state the evidence "
      "for each.",
      "Write one Point of View statement and one How Might We question for the pain point you judge most "
      "important, and justify why that HMW question is at the right altitude.",
      "Explain how you would engage residents and frontline staff to uncover the motivations behind the "
      "behaviour, and how you would handle the contradiction between what residents say and what they do."],
     ["Insight brief that combines both evidence types rather than privileging the dashboard: attendance "
      "and app data show WHAT happens; interviews and observation explain WHY. Credit explicit use of an "
      "empathy map (Says / Thinks / Does / Feels, plus Pains and Gains).",
      "The say/do contradiction is the central insight in this case — residents say they will attend and "
      "then do not. Strong answers name this gap explicitly as the place the real insight lives, rather "
      "than treating the survey responses as truth.",
      "Three acceptable pain points with evidence: (1) fear of a bad result — residents avoid information "
      "they cannot afford to act on; (2) cost uncertainty — free-text comments mention not knowing what "
      "it will cost; (3) attending alone — the nurse's observation that residents accompanied by family "
      "almost always complete, while those alone frequently leave.",
      "Other defensible pain points: language or literacy barriers at the point of registration; "
      "reluctance to enter an NRIC or personal details in a public area; travel difficulty; not "
      "understanding what the screening involves.",
      "Point of View in the required form — [user] needs [need] because [insight]. For example: 'Mr Lim, "
      "72, needs to feel accompanied through the screening because arriving alone makes an uncertain "
      "result feel unmanageable.' The insight must trace to a specific piece of evidence.",
      "How Might We at the right altitude — for example 'How might we make a first screening feel "
      "accompanied rather than solitary?'. Justification must apply the altitude test: too narrow if only "
      "one obvious solution exists (the solution is hiding in the question); too broad if no concrete "
      "answer can be pictured.",
      "Engagement approach: interviews in the residents' own context rather than in the clinic; "
      "observation of the actual arrival and registration journey; asking why repeatedly; involving "
      "frontline outreach staff as co-researchers because residents already trust them.",
      "Handling the contradiction: treat stated intention as unreliable and design around observed "
      "behaviour; ask about the last specific occasion rather than about general intentions.",
      "Credit any answer that distinguishes problem validation from solution validation and from demand "
      "validation, and notes these need three different tests."]),

    ("Task 3: Lead Agile Execution of the Programme",
     "A7",
     "Mapped to A7 — Lead design thinking projects across the organisation.",
     ["Outline how you would run delivery using agile principles — describe your board, the design "
      "thinking phases it reflects, and the fields you would track.",
      "Define the team roles, the checkpoints and the feedback loops, and explain how discovery and "
      "delivery will run in relation to one another.",
      "Convert your chosen concept into one outcome statement (a Hill) and at least three user stories "
      "with testable acceptance criteria, and explain how you would prioritise them across releases.",
      "Explain how you would lead the team to adapt and course-correct during execution — including what "
      "you would do when Finance arrives mid-sprint demanding an unplanned change."],
     ["Board design: columns reflecting flow (for example Backlog, In discovery, Ready for delivery, In "
      "progress, In review, Done), with the design thinking phases visible so the team can see which "
      "items are still being validated. Credit any board that makes work-in-progress and blockers visible "
      "— an information radiator.",
      "Roles: a Product Owner accountable for priority and the backlog (the WHAT), a Scrum Master or "
      "facilitator who removes impediments and protects focus, and a cross-functional delivery team who "
      "own the HOW. Credit naming a product trio — product, design and engineering — leading discovery.",
      "Dual-track agile explicitly described: a discovery track asking 'should we build this, and what "
      "exactly?' running simultaneously with a delivery track asking 'how do we build this well?'. They "
      "run in parallel within one team, not in sequence, and discovery feeds validated items into "
      "delivery. Discovery legitimately kills ideas.",
      "Credit identifying the mini-waterfall anti-pattern — a product manager writing requirements, "
      "handing them to a designer, who hands them to the delivery team — where each handoff looks like "
      "collaboration and is actually a queue.",
      "Hill in the required form: 'As a resident attending a first screening, I want to arrive knowing "
      "exactly what will happen and what it costs, so that I complete the screening without needing "
      "anyone to persuade me.' The 'wow' must be measurable.",
      "User stories in the form 'As a [user], I want [goal], so that [reason]', each with acceptance "
      "criteria that state a testable condition rather than restating the story. Reject criteria that "
      "merely paraphrase the story.",
      "Prioritisation across releases using the three-cake model: the Cupcake release is complete and "
      "genuinely useful on its own, not a slice of an unfinished product; Birthday Cake broadens reach; "
      "Wedding Cake is the full vision. Estimation by relative story points, with velocity counted only "
      "on fully completed work.",
      "Checkpoints and feedback loops: a daily stand-up of no more than fifteen minutes (progress, next "
      "step, blockers, no debate), a sprint review that genuinely seeks disconfirming feedback, and a "
      "retrospective that produces a small number of owned actions.",
      "The mid-sprint disruption: protect the sprint goal and route the new request into the discovery "
      "track for validation rather than swapping scope. Absorbing it teaches stakeholders that "
      "interrupting works. The request does not have to wait long — it has to be validated before it "
      "displaces committed work. Document the decision and the reasoning."]),

    ("Task 4: Develop an Organisational Strategy for Agile Design Thinking",
     "A2 · A4",
     "Mapped to A2 — Develop strategies to proliferate design thinking across the organisation, and A4 — "
     "Drive the development of new strategies to enhance products and/or services.",
     ["Propose a multi-phase strategy to scale agile design thinking from this one programme across all "
      "eleven clinics, and state what would have to change beyond running more training.",
      "Define the innovation metrics you would report to the board, covering customer, business and "
      "societal value, and justify each against a decision it would change.",
      "Explain how you would secure and sustain stakeholder buy-in, including from Finance and from the "
      "Clinic Operations team whose original proposal you did not fund.",
      "Recommend how future service innovation projects should be structured to remain responsive, and "
      "identify one systemic risk in scaling AI-assisted innovation together with the mechanism you would "
      "put in place to contain it."],
     ["Multi-phase strategy: prove it in one clinic with a real measured outcome; extend to a small "
      "number of clinics with the same method and a named metric; then institutionalise. Credit "
      "explicitly starting small with low-risk, high-value work and one to three attempts before scaling.",
      "The central point, which strong answers make explicitly: training alone produces people who know "
      "the vocabulary and cannot use it. Scaling requires changing what is MEASURED, who has AUTHORITY to "
      "decide, and what leaders visibly DO. Credit reference to DBS changing its measurement (customer "
      "hours saved), moving capability in-house, and requiring leaders to participate rather than sponsor.",
      "Also creditable: enabling toolkits and shared workspaces, a community of practice, embedded "
      "coaches rather than one-off workshops, and protected discovery time in every team.",
      "Metrics — customer value (screening completion rate, time from arrival to completion, proportion "
      "abandoning before starting), business value (cost per completed screening, downstream cost avoided "
      "through earlier detection, clinic throughput), societal value (screening uptake among "
      "under-served or non-English-speaking residents, equity of access).",
      "Every metric must pass the decision test: if a movement in the number would change no decision "
      "anyone would make, it is decoration. Explicitly reject vanity metrics such as ideas submitted, "
      "workshops run or staff trained — they always rise and prove nothing.",
      "At least one leading indicator that moves early enough to change a decision within the programme, "
      "not only a lagging outcome.",
      "Buy-in: agree the success metric before the pilot rather than after; bring evidence (a resident "
      "quote, a test result, a falsified assumption) rather than opinion; state in advance what you will "
      "stop doing if the evidence goes against you. For Clinic Operations specifically — involve them as "
      "co-owners of the discovery, show that their proposal was tested rather than dismissed, and give "
      "them the credit for any resulting change.",
      "Future structure: continuous discovery running permanently beside delivery; small cross-functional "
      "teams with real decision rights; staged funding released against validated learning; and a "
      "standing practice of killing weak initiatives to release capacity.",
      "Systemic risk in scaling AI-assisted innovation — a reinforcing loop with no balancing "
      "counterpart: schedule pressure leads to AI-generated insight being accepted without validation, "
      "which leads to decisions built on unvalidated personas, which increases rework, which increases "
      "schedule pressure. Credit the Boeing 737 MAX parallel, where every individual decision was "
      "locally rational and the system of decisions was lethal.",
      "The containment mechanism must be STRUCTURAL rather than attitudinal — a named balancing "
      "mechanism with a named owner who has the authority to stop work: for example a mandatory "
      "provenance label on every AI-generated artefact, a validation gate that a concept cannot pass "
      "without real user evidence, and a person empowered to hold the gate against schedule pressure.",
      "Credit any answer noting that the barrier is rarely detection — it is the incentive and the "
      "authority to act on what has been detected."]),
]


# ================================================================ doc helpers
def new_doc():
    d = Document()
    n = d.styles["Normal"]; n.font.name = "Arial"; n.font.size = Pt(11)
    return d


def line(d, text="", bold=False, size=11, color=DARK, after=6, align=None, italic=False):
    p = d.add_paragraph(); p.paragraph_format.space_after = Pt(after)
    if align is not None:
        p.alignment = align
    r = p.add_run(text); r.bold = bold; r.italic = italic
    r.font.size = Pt(size); r.font.color.rgb = color; r.font.name = "Arial"
    return p


def bullet(d, text, size=10.5):
    p = d.add_paragraph(style="List Bullet"); p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text); r.font.size = Pt(size); r.font.name = "Arial"; r.font.color.rgb = DARK
    return p


def cover(d, instrument):
    prodoc.add_cover_page(d, instrument, C.TITLE, C.VERSION.lstrip("v"),
                          org_logo=ORG_LOGO, course_logo=None, course_code=C.COURSE_CODE)


def candidate_block(d, instrument, duration_text):
    line(d, "A: Trainee Information", bold=True, size=12, after=6)
    line(d, "Trainee Name (as per NRIC): _______________________________", after=6)
    line(d, "Last 3 digits and alphabet of NRIC / FIN: _________________", after=6)
    line(d, "Date: __________________", after=12)
    line(d, "B: Instructions to Candidate", bold=True, size=12, after=6)
    for i, t in enumerate([
            "This is an individual exercise.",
            "This is an open book assessment. You may refer to the course slides, the Learner Guide, "
            "your activity briefs and other approved materials.",
            duration_text,
            "Write your answers in the space provided on this document.",
            f"Submit your completed answers on the LMS at {LMS_URL}",
            "All questions are open-ended. There are no multiple-choice options.",
    ], 1):
        p = d.add_paragraph(style="List Number"); p.paragraph_format.space_after = Pt(3)
        r = p.add_run(t); r.font.size = Pt(11); r.font.name = "Arial"
    line(d, "", after=6)
    line(d, "____________________________________________________________________________",
         color=GREY, after=6)
    line(d, "For Official Use Only", bold=True, after=6)
    line(d, "Grade: _____ (C / NYC)", after=6)
    line(d, "Assessor Name: _______________\t\tAssessor NRIC: _____________", after=6)
    line(d, "Date: ________________________\t\tSignature: _________________", after=6)
    d.add_page_break()


def answer_box(d, lines=10):
    """Boxed answer space — a single-cell table the candidate writes inside."""
    t = d.add_table(rows=1, cols=1); t.style = "Table Grid"
    cell = t.rows[0].cells[0]
    cell.text = ""
    for _ in range(lines):
        pp = cell.add_paragraph(); pp.paragraph_format.space_after = Pt(6)
    d.add_paragraph("")


def finish(d, path):
    prodoc.add_page_numbers(d)
    d.save(path)
    print("Saved", os.path.basename(path))


# ================================================================ build: WA paper
def build_wa_paper():
    d = new_doc()
    cover(d, "WRITTEN ASSESSMENT (SAQ)")
    candidate_block(d, "Written Assessment (SAQ)",
                    "A total of 0.5 hour is given to complete this written assessment.")
    line(d, "C: Short Answer Questions", bold=True, size=12, after=10)
    for i, (q, k, _pts, _src) in enumerate(WRITTEN, 1):
        line(d, f"Question {i}:", bold=True, size=11.5, after=4)
        line(d, f"{q}  ({k})", after=6)
        answer_box(d, lines=9)
    finish(d, os.path.join(OUT, f"WA (SAQ) - {C.SHORT_TITLE} - {VER}.docx"))


# ================================================================ build: WA answers
def build_wa_answers():
    d = new_doc()
    cover(d, "WRITTEN ASSESSMENT (SAQ) — MODEL ANSWERS")
    line(d, "Model Answers and Marking Guide", bold=True, size=14, color=BRAND, after=4,
         align=AL.CENTER)
    line(d, "Written Assessment (SAQ)  ·  6 questions  ·  0.5 hour  ·  Open book",
         size=10.5, color=GREY, after=4, align=AL.CENTER)
    line(d, "TRAINER COPY — NOT FOR RELEASE TO CANDIDATES", bold=True, size=10,
         color=RGBColor(0xDC, 0x26, 0x26), after=12, align=AL.CENTER)
    line(d, "Marking guidance: answers are suggestive and not exhaustive. Award a Competent judgement "
            "where the candidate demonstrates the underlying understanding, even if the wording or the "
            "chosen examples differ from those below. Candidates may draw on their own workplace "
            "examples.", italic=True, size=10.5, color=GREY, after=12)
    for i, (q, k, pts, src) in enumerate(WRITTEN, 1):
        line(d, f"Question {i}  ({k})", bold=True, size=12, color=BRAND, after=4)
        line(d, q, italic=True, size=10.5, color=GREY, after=6)
        line(d, "Suggestive answers (not exhaustive):", bold=True, size=10.5, after=4)
        for pt in pts:
            bullet(d, pt)
        line(d, f"Source: {src}", size=9.5, color=GREY, italic=True, after=14)
    finish(d, os.path.join(OUT, f"Answer to WA (SAQ) - {C.SHORT_TITLE} - {VER}.docx"))


# ================================================================ build: CS paper
def build_cs_paper():
    d = new_doc()
    cover(d, "CASE STUDY ASSESSMENT")
    candidate_block(d, "Case Study Assessment",
                    "A total of 1 hour is given to complete this case study assessment.")
    line(d, "C: Case Study", bold=True, size=12, after=8)
    line(d, f"Scenario: {CS_TITLE}", bold=True, size=11.5, color=BRAND, after=8)
    for para in CS_SCENARIO:
        line(d, para, after=8)
    line(d, "", after=4)
    for title, codes, mapping, reqs, _model in CS_TASKS:
        line(d, "____________________________________________________________________________",
             color=GREY, after=8)
        line(d, f"{title}   ({codes})", bold=True, size=11.5, color=BRAND, after=4)
        line(d, mapping, italic=True, size=10, color=GREY, after=8)
        line(d, "Task Requirements:", bold=True, size=11, after=4)
        for r in reqs:
            bullet(d, r, size=10.5)
        line(d, "", after=4)
        answer_box(d, lines=14)
    finish(d, os.path.join(OUT, f"CS Assessment - {C.SHORT_TITLE} - {VER}.docx"))


# ================================================================ build: CS answers
def build_cs_answers():
    d = new_doc()
    cover(d, "CASE STUDY ASSESSMENT — MODEL ANSWERS")
    line(d, "Model Answers and Marking Guide", bold=True, size=14, color=BRAND, after=4,
         align=AL.CENTER)
    line(d, "Case Study Assessment  ·  4 tasks  ·  1 hour  ·  Open book",
         size=10.5, color=GREY, after=4, align=AL.CENTER)
    line(d, "TRAINER COPY — NOT FOR RELEASE TO CANDIDATES", bold=True, size=10,
         color=RGBColor(0xDC, 0x26, 0x26), after=12, align=AL.CENTER)
    line(d, "Marking guidance: answers are suggestive and not exhaustive. The candidate is assessed on "
            "applied judgement, not on reproducing these points verbatim. Award Competent where the "
            "candidate addresses each task requirement with reasoning consistent with the course "
            "content; candidates may substitute their own workplace context.", italic=True,
         size=10.5, color=GREY, after=10)
    line(d, f"Scenario: {CS_TITLE}", bold=True, size=11.5, color=BRAND, after=4)
    line(d, "Sunrise Health Group — flat screening uptake among residents aged 60+, an operations team "
            "proposing a reminder application, and early signals pointing to fear, cost uncertainty and "
            "attending alone rather than to forgetting.", italic=True, size=10, color=GREY, after=12)
    for title, codes, mapping, reqs, model in CS_TASKS:
        line(d, f"{title}   ({codes})", bold=True, size=12, color=BRAND, after=4)
        line(d, mapping, italic=True, size=10, color=GREY, after=6)
        line(d, "Task requirements given to the candidate:", bold=True, size=10.5, after=4)
        for r in reqs:
            bullet(d, r, size=10)
        line(d, "", after=4)
        line(d, "Suggestive answers (not exhaustive):", bold=True, size=10.5, after=4)
        for m in model:
            bullet(d, m, size=10.5)
        line(d, "", after=14)
    line(d, "____________________________________________________________________________",
         color=GREY, after=6)
    line(d, "Overall judgement", bold=True, size=12, color=BRAND, after=4)
    line(d, "The candidate is assessed Competent (C) where all four tasks are attempted and the "
            "responses demonstrate: reframing before solving; synthesis of quantitative and qualitative "
            "evidence into a Point of View; agile execution with discovery and delivery running "
            "together; and a scaling strategy whose metrics would each change a decision. A candidate "
            "who proposes building the reminder application without first testing whether forgetting is "
            "the real problem has not met Task 1.", after=8)
    finish(d, os.path.join(OUT, f"Answer to CS Assessment - {C.SHORT_TITLE} - {VER}.docx"))


if __name__ == "__main__":
    os.makedirs(OUT, exist_ok=True)
    # coverage assertions — every K in the WA, every A in the CS
    ks = {k for _, k, _, _ in WRITTEN}
    assert ks >= {"K1", "K2", "K3", "K4", "K5"}, f"WA missing K codes: {ks}"
    acs = set()
    for _, codes, _, _, _ in CS_TASKS:
        for c in codes.replace("·", " ").split():
            if c.startswith("A"):
                acs.add(c)
    assert acs >= {f"A{i}" for i in range(1, 8)}, f"CS missing A codes: {sorted(acs)}"
    print(f"Coverage OK — WA covers {sorted(ks)}; CS covers {sorted(acs)}")
    build_wa_paper(); build_wa_answers(); build_cs_paper(); build_cs_answers()
    print(f"\n4 documents written to {OUT}")
